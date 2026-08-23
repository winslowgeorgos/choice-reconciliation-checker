# baas_invoice_management_page.py
import streamlit as st
import pandas as pd
from datetime import datetime, timedelta
import io
import os
import json
import uuid
import logging
import sqlite3
from pathlib import Path
import plotly.express as px
import plotly.graph_objects as go
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
import re

# Setup logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

# --- Constants ---
UPLOAD_DIR = "data/uploads"
CACHE_DIR = "data/cache"
INVOICE_DB_PATH = "data/baas_invoice.db"
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(CACHE_DIR, exist_ok=True)

# Custom CSS
CUSTOM_CSS = """
<style>
    .main-header {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 1.5rem;
        border-radius: 10px;
        margin-bottom: 2rem;
        color: white;
    }
    .stat-card {
        background: white;
        padding: 1rem;
        border-radius: 10px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        margin: 0.5rem 0;
        border-left: 4px solid #667eea;
    }
    .stButton button {
        border-radius: 8px;
        font-weight: 500;
        transition: all 0.3s ease;
    }
    .stButton button:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(0,0,0,0.15);
    }
    .invoice-preview {
        background: white;
        padding: 2rem;
        border-radius: 10px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
        margin: 1rem 0;
        border: 1px solid #e0e0e0;
    }
    .invoice-header {
        border-bottom: 2px solid #667eea;
        padding-bottom: 1rem;
        margin-bottom: 1rem;
    }
    .invoice-total {
        background: #f8f9fa;
        padding: 1rem;
        border-radius: 8px;
        margin-top: 1rem;
        text-align: right;
        font-size: 1.2rem;
        font-weight: bold;
    }
    .custom-success {
        background-color: #d4edda;
        border-left: 4px solid #28a745;
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
    }
</style>
"""

# ----------------------------- Database Manager -----------------------------
class BaaSInvoiceDB:
    """Database manager for BaaS Invoice data"""

    def __init__(self, db_path=INVOICE_DB_PATH):
        self.db_path = db_path
        self._init_database()

    def _init_database(self):
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()

        # Main invoice table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS baas_invoices (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                invoice_id TEXT UNIQUE,
                merchant_name TEXT,
                merchant_email TEXT,
                invoice_date TEXT,
                due_date TEXT,
                period TEXT,
                currency TEXT,
                platform_fee REAL,
                kyc_fee REAL,
                inbound_fee REAL,
                swift_charge REAL,
                withholding_tax REAL,
                total_amount REAL,
                status TEXT,
                payment_details TEXT,
                notes TEXT,
                created_at TEXT,
                updated_at TEXT
            )
        ''')

        # Invoice items table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS baas_invoice_items (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                invoice_id TEXT,
                description TEXT,
                quantity REAL,
                unit_price REAL,
                amount REAL,
                currency TEXT,
                FOREIGN KEY (invoice_id) REFERENCES baas_invoices(invoice_id)
            )
        ''')

        # Invoice history/log
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS baas_invoice_history (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                invoice_id TEXT,
                action TEXT,
                user TEXT,
                timestamp TEXT,
                details TEXT
            )
        ''')

        # Merchant master table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS baas_merchants (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                merchant_name TEXT UNIQUE,
                contact_email TEXT,
                contact_person TEXT,
                phone TEXT,
                billing_address TEXT,
                currency TEXT,
                vat_applicable TEXT,
                vat_rate REAL,
                status TEXT,
                created_at TEXT,
                updated_at TEXT
            )
        ''')

        conn.commit()
        conn.close()
        logger.info("BaaS Invoice database initialized")

    def save_invoice(self, invoice_data, items_data):
        """Save invoice and its items to database"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()

        # Check if invoice exists
        cursor.execute("SELECT invoice_id FROM baas_invoices WHERE invoice_id = ?", (invoice_data['invoice_id'],))
        existing = cursor.fetchone()

        if existing:
            # Update existing invoice
            cursor.execute('''
                UPDATE baas_invoices SET
                    merchant_name = ?, merchant_email = ?, invoice_date = ?,
                    due_date = ?, period = ?, currency = ?, platform_fee = ?,
                    kyc_fee = ?, inbound_fee = ?, swift_charge = ?,
                    withholding_tax = ?, total_amount = ?, status = ?,
                    payment_details = ?, notes = ?, updated_at = ?
                WHERE invoice_id = ?
            ''', (
                invoice_data['merchant_name'],
                invoice_data['merchant_email'],
                invoice_data['invoice_date'],
                invoice_data['due_date'],
                invoice_data['period'],
                invoice_data['currency'],
                invoice_data['platform_fee'],
                invoice_data['kyc_fee'],
                invoice_data['inbound_fee'],
                invoice_data['swift_charge'],
                invoice_data['withholding_tax'],
                invoice_data['total_amount'],
                invoice_data['status'],
                invoice_data['payment_details'],
                invoice_data['notes'],
                datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                invoice_data['invoice_id']
            ))
            # Delete old items
            cursor.execute("DELETE FROM baas_invoice_items WHERE invoice_id = ?", (invoice_data['invoice_id'],))
        else:
            # Insert new invoice
            cursor.execute('''
                INSERT INTO baas_invoices (
                    invoice_id, merchant_name, merchant_email, invoice_date,
                    due_date, period, currency, platform_fee, kyc_fee,
                    inbound_fee, swift_charge, withholding_tax, total_amount,
                    status, payment_details, notes, created_at, updated_at
                ) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
            ''', (
                invoice_data['invoice_id'],
                invoice_data['merchant_name'],
                invoice_data['merchant_email'],
                invoice_data['invoice_date'],
                invoice_data['due_date'],
                invoice_data['period'],
                invoice_data['currency'],
                invoice_data['platform_fee'],
                invoice_data['kyc_fee'],
                invoice_data['inbound_fee'],
                invoice_data['swift_charge'],
                invoice_data['withholding_tax'],
                invoice_data['total_amount'],
                invoice_data['status'],
                invoice_data['payment_details'],
                invoice_data['notes'],
                datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            ))

        # Insert items
        for item in items_data:
            cursor.execute('''
                INSERT INTO baas_invoice_items (
                    invoice_id, description, quantity, unit_price, amount, currency
                ) VALUES (?,?,?,?,?,?)
            ''', (
                invoice_data['invoice_id'],
                item['description'],
                item['quantity'],
                item['unit_price'],
                item['amount'],
                item['currency']
            ))

        conn.commit()
        conn.close()
        logger.info(f"Saved invoice {invoice_data['invoice_id']}")

    def get_invoices(self, status=None):
        """Get all invoices, optionally filtered by status"""
        conn = sqlite3.connect(self.db_path)
        query = "SELECT * FROM baas_invoices"
        if status:
            query += " WHERE status = ?"
            df = pd.read_sql_query(query, conn, params=(status,))
        else:
            df = pd.read_sql_query(query, conn)
        conn.close()
        return df

    def get_invoice_by_id(self, invoice_id):
        """Get a single invoice by ID"""
        conn = sqlite3.connect(self.db_path)
        df = pd.read_sql_query("SELECT * FROM baas_invoices WHERE invoice_id = ?", conn, params=(invoice_id,))
        conn.close()
        if df.empty:
            return None
        return df.iloc[0].to_dict()

    def get_invoice_items(self, invoice_id):
        """Get items for a specific invoice"""
        conn = sqlite3.connect(self.db_path)
        df = pd.read_sql_query("SELECT * FROM baas_invoice_items WHERE invoice_id = ?", conn, params=(invoice_id,))
        conn.close()
        return df

    def update_invoice_status(self, invoice_id, status):
        """Update invoice status"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute("UPDATE baas_invoices SET status = ?, updated_at = ? WHERE invoice_id = ?",
                      (status, datetime.now().strftime('%Y-%m-%d %H:%M:%S'), invoice_id))
        conn.commit()
        conn.close()

    def log_history(self, invoice_id, action, user, details=""):
        """Log invoice history"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO baas_invoice_history (invoice_id, action, user, timestamp, details)
            VALUES (?,?,?,?,?)
        ''', (invoice_id, action, user, datetime.now().strftime('%Y-%m-%d %H:%M:%S'), details))
        conn.commit()
        conn.close()

    def get_available_dates(self):
        """Get all available invoice dates"""
        conn = sqlite3.connect(self.db_path)
        df = pd.read_sql_query("SELECT DISTINCT invoice_date FROM baas_invoices ORDER BY invoice_date DESC", conn)
        conn.close()
        return df['invoice_date'].tolist() if not df.empty else []


# Initialize DB
invoice_db = BaaSInvoiceDB()

# ----------------------------- Helper Functions -----------------------------
def parse_baas_income_description(description):
    """
    Parse the BaaS income description to extract invoice components.
    
    Example: "HONGKONG FORTUNETECH BAAS SERVICE FEE FOR JULY 2026 (PLATFORM FEE - 1500usd*129.29; Inbound Fee - 10.56 usd*129.29; Inbound Swift Charge - 7usd*129.29)"
    
    Returns: {
        'merchant_name': 'HONGKONG FORTUNETECH',
        'period': 'JULY 2026',
        'components': [
            {'type': 'platform_fee', 'amount': 1500, 'currency': 'USD', 'rate': 129.29},
            {'type': 'inbound_fee', 'amount': 10.56, 'currency': 'USD', 'rate': 129.29},
            {'type': 'swift_charge', 'amount': 7, 'currency': 'USD', 'rate': 129.29}
        ]
    }
    """
    result = {
        'merchant_name': '',
        'period': '',
        'components': [],
        'currency': 'KES',
        'withholding_tax': 0
    }

    if not description or pd.isna(description):
        return result

    desc_str = str(description).strip()
    
    # Extract merchant name (before "BAAS SERVICE FEE" or "PLATFORM FEE")
    merchant_match = re.search(r'^(.*?)\s+(?:BAAS SERVICE FEE|PLATFORM FEE|SERVICE FEE)', desc_str, re.IGNORECASE)
    if merchant_match:
        result['merchant_name'] = merchant_match.group(1).strip()
    else:
        # Try alternative pattern
        merchant_match = re.search(r'^(.*?)\s+FOR\s+', desc_str, re.IGNORECASE)
        if merchant_match:
            result['merchant_name'] = merchant_match.group(1).strip()

    # Extract period (e.g., "JULY 2026" or "JULY 2026 (PLATFORM...")
    period_match = re.search(r'FOR\s+([A-Z]+\s+\d{4})', desc_str, re.IGNORECASE)
    if period_match:
        result['period'] = period_match.group(1)

    # Extract components from parentheses
    paren_match = re.search(r'\((.*?)\)', desc_str)
    if paren_match:
        components_str = paren_match.group(1)
        # Split by semicolon
        parts = components_str.split(';')
        for part in parts:
            part = part.strip()
            # Check for different component types
            if 'PLATFORM FEE' in part.upper():
                # Extract amount and currency
                amount_match = re.search(r'([\d,]+\.?\d*)\s*([A-Z]{3})?', part)
                if amount_match:
                    amount = float(amount_match.group(1).replace(',', ''))
                    currency = amount_match.group(2) if amount_match.group(2) else 'USD'
                    result['components'].append({
                        'type': 'platform_fee',
                        'amount': amount,
                        'currency': currency,
                        'rate': None
                    })
            elif 'INBOUND FEE' in part.upper() or 'INBOUND TXN FEE' in part.upper():
                amount_match = re.search(r'([\d,]+\.?\d*)\s*([A-Z]{3})?', part)
                if amount_match:
                    amount = float(amount_match.group(1).replace(',', ''))
                    currency = amount_match.group(2) if amount_match.group(2) else 'USD'
                    result['components'].append({
                        'type': 'inbound_fee',
                        'amount': amount,
                        'currency': currency,
                        'rate': None
                    })
            elif 'SWIFT CHARGE' in part.upper() or 'INCOMING SWIFT' in part.upper():
                amount_match = re.search(r'([\d,]+\.?\d*)\s*([A-Z]{3})?', part)
                if amount_match:
                    amount = float(amount_match.group(1).replace(',', ''))
                    currency = amount_match.group(2) if amount_match.group(2) else 'USD'
                    result['components'].append({
                        'type': 'swift_charge',
                        'amount': amount,
                        'currency': currency,
                        'rate': None
                    })
            elif 'KYC' in part.upper() or 'KYB' in part.upper():
                amount_match = re.search(r'([\d,]+\.?\d*)\s*([A-Z]{3})?', part)
                if amount_match:
                    amount = float(amount_match.group(1).replace(',', ''))
                    currency = amount_match.group(2) if amount_match.group(2) else 'KES'
                    result['components'].append({
                        'type': 'kyc_fee',
                        'amount': amount,
                        'currency': currency,
                        'rate': None
                    })

    return result

def parse_baas_row(row):
    """
    Parse a BaaS income row from the Excel file.
    Expected columns: Client Name, Invoice Sent, Amount Received, etc.
    """
    result = {
        'merchant_name': '',
        'period': '',
        'platform_fee': 0,
        'kyc_fee': 0,
        'inbound_fee': 0,
        'swift_charge': 0,
        'currency': 'KES',
        'description': '',
        'invoice_sent': '',
        'amount_received': ''
    }

    # Try to get merchant name from index or column
    if 'Client Name' in row.index:
        result['merchant_name'] = row.get('Client Name', '')
    elif 'Merchant' in row.index:
        result['merchant_name'] = row.get('Merchant', '')

    # Get description from various possible column names
    for col in ['Description', 'Details', 'Invoice Description']:
        if col in row.index:
            result['description'] = row.get(col, '')
            break

    # Parse the description to get components
    if result['description']:
        parsed = parse_baas_income_description(result['description'])
        result['merchant_name'] = parsed['merchant_name'] or result['merchant_name']
        result['period'] = parsed['period']
        result['currency'] = parsed['currency']
        
        for comp in parsed['components']:
            if comp['type'] == 'platform_fee':
                result['platform_fee'] = comp['amount']
            elif comp['type'] == 'kyc_fee':
                result['kyc_fee'] = comp['amount']
            elif comp['type'] == 'inbound_fee':
                result['inbound_fee'] = comp['amount']
            elif comp['type'] == 'swift_charge':
                result['swift_charge'] = comp['amount']

    # Get invoice status
    for col in ['Invoice Sent', 'Status', 'Sent']:
        if col in row.index:
            result['invoice_sent'] = row.get(col, '')
            break

    for col in ['Amount Received', 'Received', 'Paid']:
        if col in row.index:
            result['amount_received'] = row.get(col, '')
            break

    return result

def generate_invoice_id(merchant_name, period):
    """Generate a unique invoice ID"""
    prefix = ''.join([word[0] for word in merchant_name.split() if word]).upper()
    if not prefix:
        prefix = 'BaaS'
    period_clean = period.replace(' ', '').upper()
    timestamp = datetime.now().strftime('%Y%m%d')
    return f"{prefix}/{timestamp}/{period_clean}"

def calculate_withholding_tax(platform_fee, currency='KES'):
    """
    Calculate withholding tax (5% of platform fee)
    Only applies to KES transactions
    """
    if currency.upper() == 'KES' and platform_fee > 0:
        return platform_fee * 0.05
    return 0

def get_invoice_exchange_rate(currency):
    """Get exchange rate for currency to KES"""
    # This should use live rates, but fallback to defaults
    rates = {
        'USD': 129.29,
        'EUR': 150.64,
        'GBP': 173.38,
        'CNY': 18.15,
        'KES': 1.0
    }
    return rates.get(currency.upper(), 129.29)

def convert_currency_to_kes(amount, from_currency):
    """Convert amount from given currency to KES"""
    if from_currency.upper() == 'KES':
        return amount
    rate = get_invoice_exchange_rate(from_currency)
    return amount * rate

def create_invoice_document(invoice_data, items_data, merchant_data=None):
    """
    Create a Word document invoice from the data
    """
    doc = Document()
    
    # Set up document margins
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # === HEADER ===
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Company logo/title
    title = header.add_run("CHOICE MICROFINANCE BANK LIMITED")
    title.bold = True
    title.font.size = Pt(16)
    title.font.color.rgb = RGBColor(75, 45, 143)
    
    doc.add_paragraph("BaaS Service Fee Invoice")
    doc.add_paragraph("-" * 50)
    doc.add_paragraph()

    # === INVOICE DETAILS ===
    # Left: To, Right: Invoice details
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set column widths
    table.columns[0].width = Inches(3)
    table.columns[1].width = Inches(3)
    
    # Row 1: To
    cell_to = table.cell(0, 0)
    cell_to.text = "To:"
    cell_to.paragraphs[0].runs[0].bold = True
    
    cell_name = table.cell(0, 1)
    cell_name.text = invoice_data.get('merchant_name', '')
    cell_name.paragraphs[0].runs[0].bold = True
    
    # Row 2: Email
    cell_email_label = table.cell(1, 0)
    cell_email_label.text = "Email:"
    cell_email_label.paragraphs[0].runs[0].bold = True
    
    cell_email = table.cell(1, 1)
    cell_email.text = invoice_data.get('merchant_email', '')
    
    # Row 3: Invoice No
    cell_inv_label = table.cell(2, 0)
    cell_inv_label.text = "Invoice No:"
    cell_inv_label.paragraphs[0].runs[0].bold = True
    
    cell_inv = table.cell(2, 1)
    cell_inv.text = invoice_data.get('invoice_id', '')
    
    # Row 4: Date
    cell_date_label = table.cell(3, 0)
    cell_date_label.text = "Date:"
    cell_date_label.paragraphs[0].runs[0].bold = True
    
    cell_date = table.cell(3, 1)
    cell_date.text = invoice_data.get('invoice_date', '')
    
    # Row 5: Due Date
    cell_due_label = table.cell(4, 0)
    cell_due_label.text = "Due Date:"
    cell_due_label.paragraphs[0].runs[0].bold = True
    
    cell_due = table.cell(4, 1)
    cell_due.text = invoice_data.get('due_date', '')
    
    # Row 6: Period
    cell_period_label = table.cell(5, 0)
    cell_period_label.text = "Period:"
    cell_period_label.paragraphs[0].runs[0].bold = True
    
    cell_period = table.cell(5, 1)
    cell_period.text = invoice_data.get('period', '')
    
    # Row 7: Currency
    cell_currency_label = table.cell(6, 0)
    cell_currency_label.text = "Currency:"
    cell_currency_label.paragraphs[0].runs[0].bold = True
    
    cell_currency = table.cell(6, 1)
    cell_currency.text = invoice_data.get('currency', 'KES')
    
    # Row 8: Status
    cell_status_label = table.cell(7, 0)
    cell_status_label.text = "Status:"
    cell_status_label.paragraphs[0].runs[0].bold = True
    
    cell_status = table.cell(7, 1)
    cell_status.text = invoice_data.get('status', 'Draft')

    doc.add_paragraph()

    # === ITEMS TABLE ===
    doc.add_heading("Description", level=2)
    
    # Create items table
    item_table = doc.add_table(rows=1 + len(items_data), cols=4)
    item_table.style = 'Table Grid'
    item_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set column widths
    for i in range(4):
        item_table.columns[i].width = Inches(1.5)
    
    # Header row
    header_cells = item_table.rows[0].cells
    header_cells[0].text = "DESCRIPTION"
    header_cells[1].text = "QUANTITY"
    header_cells[2].text = "UNIT PRICE"
    header_cells[3].text = "AMOUNT"
    
    for cell in header_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Data rows
    total_amount = 0
    for i, item in enumerate(items_data):
        row = item_table.rows[i + 1]
        row.cells[0].text = item.get('description', '')
        row.cells[1].text = str(item.get('quantity', 1))
        row.cells[2].text = f"{item.get('unit_price', 0):,.2f}"
        row.cells[3].text = f"{item.get('amount', 0):,.2f}"
        
        # Right align amount columns
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        total_amount += item.get('amount', 0)

    doc.add_paragraph()

    # === TOTALS ===
    # Subtotal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Subtotal: ").bold = True
    p.add_run(f"{invoice_data.get('currency', 'KES')} {total_amount:,.2f}")

    # Withholding Tax (if applicable)
    withholding_tax = invoice_data.get('withholding_tax', 0)
    if withholding_tax > 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run("Less: Withholding Tax (5%): ").bold = True
        p.add_run(f"{invoice_data.get('currency', 'KES')} {withholding_tax:,.2f}")

    # Total
    net_total = total_amount - withholding_tax
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("TOTAL NET PAYABLE: ").bold = True
    p.add_run(f"{invoice_data.get('currency', 'KES')} {net_total:,.2f}")

    doc.add_paragraph()

    # === PAYMENT DETAILS ===
    doc.add_heading("Payment Details:", level=2)
    
    payment_details = invoice_data.get('payment_details', '')
    if payment_details:
        lines = payment_details.split('\n')
        for line in lines:
            p = doc.add_paragraph(line)
            p.paragraph_format.left_indent = Inches(0.5)

    # === NOTES ===
    if invoice_data.get('notes'):
        doc.add_heading("Notes:", level=2)
        p = doc.add_paragraph(invoice_data.get('notes', ''))
        p.paragraph_format.left_indent = Inches(0.5)

    return doc

def send_invoice_email(invoice_data, pdf_buffer, to_email):
    """
    Send invoice via email
    """
    # This is a placeholder - implement actual email sending
    # For production, use proper SMTP configuration
    try:
        # SMTP configuration (use environment variables in production)
        smtp_server = "smtp.gmail.com"
        smtp_port = 587
        sender_email = "your-email@example.com"
        sender_password = "your-password"
        
        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = to_email
        msg['Subject'] = f"Invoice {invoice_data.get('invoice_id', '')}"
        
        body = f"""
        Dear {invoice_data.get('merchant_name', '')},
        
        Please find attached invoice {invoice_data.get('invoice_id', '')} for BaaS service fee.
        
        Amount: {invoice_data.get('currency', 'KES')} {invoice_data.get('total_amount', 0):,.2f}
        Due Date: {invoice_data.get('due_date', '')}
        
        Please remit payment by the due date.
        
        Best regards,
        Choice Microfinance Bank Limited
        """
        
        msg.attach(MIMEText(body, 'plain'))
        
        # Attach document
        part = MIMEBase('application', 'octet-stream')
        part.set_payload(pdf_buffer.getvalue())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', f'attachment; filename= invoice_{invoice_data.get("invoice_id", "")}.docx')
        msg.attach(part)
        
        # Send email
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()
        server.login(sender_email, sender_password)
        server.send_message(msg)
        server.quit()
        
        return True, "Email sent successfully!"
    except Exception as e:
        return False, f"Failed to send email: {str(e)}"

# ----------------------------- Data Processing Functions -----------------------------
def process_baas_excel_file(uploaded_file):
    """
    Process the BaaS Income Excel file and extract invoice data
    """
    try:
        # Load the Excel file
        xls = pd.ExcelFile(uploaded_file)
        
        # Check for the GL Breakdown sheet (latest month)
        sheets = xls.sheet_names
        
        # Find the "GL Breakdown" sheet for the current month
        gl_sheets = [s for s in sheets if 'GL Breakdown' in s and not 'GL Breakdown Jun' in s and not 'GL Breakdown Jul' in s]
        
        if not gl_sheets:
            st.error("Could not find GL Breakdown sheet in the file.")
            return None
        
        # Use the most recent GL Breakdown sheet
        gl_sheet = gl_sheets[0]
        df = pd.read_excel(uploaded_file, sheet_name=gl_sheet)
        
        # Parse the data
        invoices = []
        
        # Find the relevant rows (non-empty, containing BAAS service fee)
        for idx, row in df.iterrows():
            # Skip empty rows
            if row.isna().all():
                continue
            
            # Look for BAAS service fee descriptions
            desc_cols = [col for col in df.columns if any(keyword in str(col).lower() for keyword in ['description', 'client', 'merchant', 'baas'])]
            
            if not desc_cols:
                # Try to find by looking for specific text patterns
                for col in df.columns:
                    val = str(row.get(col, '')).strip()
                    if 'BAAS SERVICE FEE' in val.upper() or 'PLATFORM FEE' in val.upper():
                        desc = val
                        break
                else:
                    continue
            else:
                desc_col = desc_cols[0]
                desc = str(row.get(desc_col, '')).strip()
                
                # Skip if it's a header or empty
                if not desc or desc.upper() in ['BAAS INCOME', 'GL', 'TRACKER', 'DIFFERENCE', 'EXPLANATION']:
                    continue
            
            # Parse the description
            parsed = parse_baas_income_description(desc)
            
            if parsed['merchant_name']:
                # Get amounts from the row
                platform_fee = parsed.get('platform_fee', 0)
                kyc_fee = parsed.get('kyc_fee', 0)
                inbound_fee = parsed.get('inbound_fee', 0)
                swift_charge = parsed.get('swift_charge', 0)
                
                # If we couldn't parse amounts from description, try to get them from columns
                if platform_fee == 0:
                    for col in ['Platform Fee', 'GL', 'Amount']:
                        if col in df.columns:
                            val = safe_float(row.get(col))
                            if val and val > 0:
                                platform_fee = val
                                break
                
                # Determine currency
                currency = 'KES'
                if 'USD' in desc or 'usd' in desc:
                    currency = 'USD'
                elif 'EUR' in desc or 'eur' in desc:
                    currency = 'EUR'
                elif 'GBP' in desc or 'gbp' in desc:
                    currency = 'GBP'
                
                # Calculate withholding tax (5% of platform fee for KES)
                withholding_tax = calculate_withholding_tax(platform_fee, currency)
                
                # Calculate total
                total_amount = platform_fee + kyc_fee + inbound_fee + swift_charge - withholding_tax
                
                # Create invoice data
                invoice_data = {
                    'merchant_name': parsed['merchant_name'],
                    'period': parsed['period'],
                    'currency': currency,
                    'platform_fee': platform_fee,
                    'kyc_fee': kyc_fee,
                    'inbound_fee': inbound_fee,
                    'swift_charge': swift_charge,
                    'withholding_tax': withholding_tax,
                    'total_amount': total_amount,
                    'description': desc,
                    'status': 'Draft'
                }
                
                # Get invoice status from row
                for col in ['Status', 'Invoice Sent', 'Sent']:
                    if col in df.columns:
                        val = str(row.get(col, '')).strip().lower()
                        if val in ['sent', 'paid']:
                            invoice_data['status'] = 'Sent' if val == 'sent' else 'Paid'
                
                invoices.append(invoice_data)
        
        return invoices
        
    except Exception as e:
        st.error(f"Error processing file: {str(e)}")
        return None

def safe_float(x):
    if pd.isna(x) or x is None:
        return None
    try:
        cleaned_x = str(x).replace(',', '').strip()
        return float(cleaned_x)
    except (ValueError, TypeError):
        return None

# ----------------------------- Main App Function -----------------------------
def baas_invoice_management_app():
    st.set_page_config(layout="wide", page_title="BaaS Invoice Management")
    st.markdown(CUSTOM_CSS, unsafe_allow_html=True)
    
    # Initialize session state
    if 'baas_invoice_initialized' not in st.session_state:
        st.session_state.baas_invoices = []
        st.session_state.baas_invoice_initialized = True
    
    st.markdown("""
    <div class="main-header">
        <h1>📄 BaaS Invoice Management</h1>
        <p>Generate, preview, download, and send invoices from BaaS Income data</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Sidebar navigation
    with st.sidebar:
        st.markdown("### 📋 Invoice Management")
        module_option = st.radio(
            "Select Action",
            [
                "📤 Import Data",
                "📋 View Invoices",
                "✏️ Create Invoice",
                "📊 Analytics"
            ],
            key="baas_invoice_module"
        )
        
        st.markdown("---")
        st.markdown("### 💡 Quick Actions")
        if st.button("🔄 Refresh Data", use_container_width=True):
            st.rerun()
        
        if st.button("🗑️ Reset All", use_container_width=True):
            if st.session_state.baas_invoices:
                st.session_state.baas_invoices = []
                st.success("All invoices cleared!")
            st.rerun()
    
    # Main content
    if module_option == "📤 Import Data":
        st.markdown("### 📤 Import BaaS Income Data")
        st.info("Upload the BaaS Income Final Excel file to automatically extract invoice data.")
        
        uploaded_file = st.file_uploader(
            "Upload BaaS Income Excel File",
            type=["xlsx", "xls"],
            key="baas_upload"
        )
        
        if uploaded_file:
            with st.spinner("Processing file..."):
                invoices = process_baas_excel_file(uploaded_file)
                
                if invoices:
                    st.session_state.baas_invoices = invoices
                    st.success(f"✅ Successfully extracted {len(invoices)} invoices from the file!")
                    
                    # Preview extracted data
                    with st.expander("📊 Extracted Invoice Data", expanded=True):
                        preview_df = pd.DataFrame(invoices)
                        preview_df = preview_df[[
                            'merchant_name', 'period', 'currency',
                            'platform_fee', 'kyc_fee', 'inbound_fee',
                            'swift_charge', 'withholding_tax', 'total_amount',
                            'status'
                        ]]
                        st.dataframe(preview_df, use_container_width=True)
                    
                    # Option to save all
                    if st.button("💾 Save All Invoices to Database", type="primary"):
                        saved_count = 0
                        for invoice in invoices:
                            invoice_id = generate_invoice_id(invoice['merchant_name'], invoice['period'])
                            invoice_data = {
                                'invoice_id': invoice_id,
                                'merchant_name': invoice['merchant_name'],
                                'merchant_email': '',
                                'invoice_date': datetime.now().strftime('%d/%m/%Y'),
                                'due_date': (datetime.now() + timedelta(days=7)).strftime('%d/%m/%Y'),
                                'period': invoice['period'],
                                'currency': invoice['currency'],
                                'platform_fee': invoice['platform_fee'],
                                'kyc_fee': invoice['kyc_fee'],
                                'inbound_fee': invoice['inbound_fee'],
                                'swift_charge': invoice['swift_charge'],
                                'withholding_tax': invoice['withholding_tax'],
                                'total_amount': invoice['total_amount'],
                                'status': invoice['status'],
                                'payment_details': 'Please see attached invoice for payment details.',
                                'notes': f'BaaS Service Fee for {invoice["period"]}'
                            }
                            
                            # Create items
                            items = []
                            if invoice['platform_fee'] > 0:
                                items.append({
                                    'description': f"Monthly Platform Fee ({invoice['period']})",
                                    'quantity': 1,
                                    'unit_price': invoice['platform_fee'],
                                    'amount': invoice['platform_fee'],
                                    'currency': invoice['currency']
                                })
                            if invoice['kyc_fee'] > 0:
                                items.append({
                                    'description': f"KYC/KYB Fee ({invoice['period']})",
                                    'quantity': 1,
                                    'unit_price': invoice['kyc_fee'],
                                    'amount': invoice['kyc_fee'],
                                    'currency': invoice['currency']
                                })
                            if invoice['inbound_fee'] > 0:
                                items.append({
                                    'description': f"Inbound Transaction Fee ({invoice['period']})",
                                    'quantity': 1,
                                    'unit_price': invoice['inbound_fee'],
                                    'amount': invoice['inbound_fee'],
                                    'currency': invoice['currency']
                                })
                            if invoice['swift_charge'] > 0:
                                items.append({
                                    'description': f"Incoming Swift Charges ({invoice['period']})",
                                    'quantity': 1,
                                    'unit_price': invoice['swift_charge'],
                                    'amount': invoice['swift_charge'],
                                    'currency': invoice['currency']
                                })
                            
                            invoice_db.save_invoice(invoice_data, items)
                            invoice_db.log_history(invoice_id, 'created', 'system', 'Invoice created from import')
                            saved_count += 1
                        
                        st.success(f"✅ Saved {saved_count} invoices to database!")
    
    elif module_option == "📋 View Invoices":
        st.markdown("### 📋 Invoice List")
        
        # Filters
        col1, col2, col3 = st.columns(3)
        with col1:
            status_filter = st.selectbox(
                "Filter by Status",
                ["All", "Draft", "Sent", "Paid"]
            )
        with col2:
            currency_filter = st.selectbox(
                "Filter by Currency",
                ["All", "KES", "USD", "EUR", "GBP"]
            )
        with col3:
            search_term = st.text_input("Search by Merchant Name", "")
        
        # Load invoices
        if status_filter != "All":
            df = invoice_db.get_invoices(status=status_filter)
        else:
            df = invoice_db.get_invoices()
        
        if not df.empty:
            # Apply filters
            if currency_filter != "All":
                df = df[df['currency'] == currency_filter]
            if search_term:
                df = df[df['merchant_name'].str.contains(search_term, case=False, na=False)]
            
            # Display invoices
            st.dataframe(
                df[[
                    'invoice_id', 'merchant_name', 'period',
                    'currency', 'total_amount', 'status',
                    'invoice_date', 'due_date'
                ]],
                use_container_width=True
            )
            
            # Action buttons for selected invoice
            st.markdown("---")
            st.markdown("### 📄 Invoice Actions")
            
            selected_invoice_id = st.selectbox(
                "Select an invoice to manage:",
                df['invoice_id'].tolist() if not df.empty else []
            )
            
            if selected_invoice_id:
                invoice_data = invoice_db.get_invoice_by_id(selected_invoice_id)
                items_df = invoice_db.get_invoice_items(selected_invoice_id)
                
                if invoice_data:
                    col1, col2, col3, col4, col5 = st.columns(5)
                    
                    with col1:
                        if st.button("👁️ Preview", use_container_width=True):
                            st.session_state.preview_invoice_id = selected_invoice_id
                            st.rerun()
                    
                    with col2:
                        if st.button("📥 Download", use_container_width=True):
                            # Generate and download invoice
                            doc = create_invoice_document(invoice_data, items_df.to_dict('records'))
                            doc_buffer = BytesIO()
                            doc.save(doc_buffer)
                            doc_buffer.seek(0)
                            
                            st.download_button(
                                label="📥 Download Invoice",
                                data=doc_buffer,
                                file_name=f"invoice_{selected_invoice_id}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                    
                    with col3:
                        if st.button("📧 Send Email", use_container_width=True):
                            # Show email form
                            st.session_state.send_invoice_id = selected_invoice_id
                            st.rerun()
                    
                    with col4:
                        if invoice_data['status'] != 'Sent':
                            if st.button("📨 Mark as Sent", use_container_width=True):
                                invoice_db.update_invoice_status(selected_invoice_id, 'Sent')
                                invoice_db.log_history(selected_invoice_id, 'status_change', 'user', 'Marked as Sent')
                                st.success("Invoice marked as Sent")
                                st.rerun()
                    
                    with col5:
                        if invoice_data['status'] != 'Paid':
                            if st.button("✅ Mark as Paid", use_container_width=True):
                                invoice_db.update_invoice_status(selected_invoice_id, 'Paid')
                                invoice_db.log_history(selected_invoice_id, 'status_change', 'user', 'Marked as Paid')
                                st.success("Invoice marked as Paid")
                                st.rerun()
            
            # Preview invoice
            if hasattr(st.session_state, 'preview_invoice_id') and st.session_state.preview_invoice_id:
                invoice_data = invoice_db.get_invoice_by_id(st.session_state.preview_invoice_id)
                items_df = invoice_db.get_invoice_items(st.session_state.preview_invoice_id)
                
                if invoice_data:
                    st.markdown("---")
                    st.markdown("### 📄 Invoice Preview")
                    
                    # Display invoice in a styled container
                    with st.container():
                        st.markdown('<div class="invoice-preview">', unsafe_allow_html=True)
                        
                        # Header
                        st.markdown("""
                        <div class="invoice-header" style="text-align:center;">
                            <h2 style="color:#4B2D8F;">CHOICE MICROFINANCE BANK LIMITED</h2>
                            <p><strong>BaaS Service Fee Invoice</strong></p>
                            <hr>
                        </div>
                        """, unsafe_allow_html=True)
                        
                        # Invoice details
                        col1, col2 = st.columns(2)
                        with col1:
                            st.markdown(f"**To:** {invoice_data.get('merchant_name', '')}")
                            st.markdown(f"**Email:** {invoice_data.get('merchant_email', '')}")
                        with col2:
                            st.markdown(f"**Invoice No:** {invoice_data.get('invoice_id', '')}")
                            st.markdown(f"**Date:** {invoice_data.get('invoice_date', '')}")
                            st.markdown(f"**Due Date:** {invoice_data.get('due_date', '')}")
                            st.markdown(f"**Period:** {invoice_data.get('period', '')}")
                            st.markdown(f"**Currency:** {invoice_data.get('currency', 'KES')}")
                        
                        st.markdown("---")
                        
                        # Items
                        st.markdown("### Description")
                        items = items_df.to_dict('records')
                        total = 0
                        for item in items:
                            col1, col2, col3, col4 = st.columns([3,1,1,1])
                            with col1:
                                st.write(item['description'])
                            with col2:
                                st.write(f"{item['quantity']:.0f}")
                            with col3:
                                st.write(f"{item['unit_price']:,.2f}")
                            with col4:
                                st.write(f"{item['amount']:,.2f}")
                            total += item['amount']
                        
                        st.markdown("---")
                        
                        # Totals
                        col1, col2 = st.columns([2,1])
                        with col2:
                            st.markdown(f"**Subtotal:** {invoice_data.get('currency', 'KES')} {total:,.2f}")
                            if invoice_data.get('withholding_tax', 0) > 0:
                                st.markdown(f"**Less: Withholding Tax (5%):** {invoice_data.get('currency', 'KES')} {invoice_data.get('withholding_tax', 0):,.2f}")
                            st.markdown(f"**TOTAL NET PAYABLE:** {invoice_data.get('currency', 'KES')} {invoice_data.get('total_amount', 0):,.2f}")
                        
                        st.markdown("---")
                        
                        # Payment details
                        st.markdown("### Payment Details:")
                        st.text(invoice_data.get('payment_details', 'Please see attached invoice for payment details.'))
                        
                        st.markdown('</div>', unsafe_allow_html=True)
                    
                    # Clear preview state
                    if st.button("Close Preview"):
                        del st.session_state.preview_invoice_id
                        st.rerun()
            
            # Send email form
            if hasattr(st.session_state, 'send_invoice_id') and st.session_state.send_invoice_id:
                invoice_data = invoice_db.get_invoice_by_id(st.session_state.send_invoice_id)
                
                if invoice_data:
                    st.markdown("---")
                    st.markdown("### 📧 Send Invoice via Email")
                    
                    with st.form("send_invoice_email_form"):
                        to_email = st.text_input("Recipient Email", value=invoice_data.get('merchant_email', ''))
                        cc_email = st.text_input("CC Email (optional)")
                        subject = st.text_input("Subject", value=f"Invoice {invoice_data.get('invoice_id', '')}")
                        
                        message = st.text_area(
                            "Message",
                            value=f"""Dear {invoice_data.get('merchant_name', '')},
                    
                    Please find attached invoice {invoice_data.get('invoice_id', '')} for BaaS service fee.
                    
                    Amount: {invoice_data.get('currency', 'KES')} {invoice_data.get('total_amount', 0):,.2f}
                    Due Date: {invoice_data.get('due_date', '')}
                    
                    Please remit payment by the due date.
                    
                    Best regards,
                    Choice Microfinance Bank Limited"""
                        )
                        
                        col1, col2 = st.columns(2)
                        with col1:
                            send_button = st.form_submit_button("📧 Send Email", type="primary", use_container_width=True)
                        with col2:
                            cancel_button = st.form_submit_button("Cancel", use_container_width=True)
                        
                        if send_button and to_email:
                            # Generate document for attachment
                            items_df = invoice_db.get_invoice_items(st.session_state.send_invoice_id)
                            doc = create_invoice_document(invoice_data, items_df.to_dict('records'))
                            doc_buffer = BytesIO()
                            doc.save(doc_buffer)
                            doc_buffer.seek(0)
                            
                            # Send email
                            success, message_result = send_invoice_email(
                                invoice_data,
                                doc_buffer,
                                to_email
                            )
                            
                            if success:
                                st.success("✅ Email sent successfully!")
                                invoice_db.log_history(
                                    invoice_data['invoice_id'],
                                    'email_sent',
                                    'user',
                                    f"Email sent to {to_email}"
                                )
                                st.session_state.send_invoice_id = None
                                st.rerun()
                            else:
                                st.error(f"❌ Failed to send email: {message_result}")
                        
                        if cancel_button:
                            st.session_state.send_invoice_id = None
                            st.rerun()
        else:
            st.info("No invoices found. Import data first or create a new invoice.")
    
    elif module_option == "✏️ Create Invoice":
        st.markdown("### ✏️ Create New Invoice")
        
        col1, col2 = st.columns(2)
        with col1:
            merchant_name = st.text_input("Merchant Name *", key="inv_merchant")
            period = st.text_input("Invoice Period *", placeholder="e.g., JULY 2026", key="inv_period")
            currency = st.selectbox("Currency", ["KES", "USD", "EUR", "GBP"], key="inv_currency")
            status = st.selectbox("Status", ["Draft", "Sent", "Paid"], key="inv_status")
        
        with col2:
            invoice_date = st.date_input("Invoice Date", datetime.now(), key="inv_date")
            due_date = st.date_input("Due Date", datetime.now() + timedelta(days=7), key="inv_due")
            merchant_email = st.text_input("Merchant Email", placeholder="example@domain.com", key="inv_email")
        
        st.markdown("---")
        st.markdown("### 💰 Invoice Items")
        
        # Add items dynamically
        if 'inv_items' not in st.session_state:
            st.session_state.inv_items = []
        
        # Show existing items
        for i, item in enumerate(st.session_state.inv_items):
            col1, col2, col3, col4, col5 = st.columns([3,1,1,1,0.5])
            with col1:
                st.text_input("Description", value=item.get('description', ''), key=f"inv_item_desc_{i}")
            with col2:
                st.number_input("Qty", value=item.get('quantity', 1), step=1, key=f"inv_item_qty_{i}")
            with col3:
                st.number_input("Unit Price", value=item.get('unit_price', 0.0), step=0.01, key=f"inv_item_price_{i}")
            with col4:
                st.text(f"Amount: {item.get('amount', 0):.2f}")
            with col5:
                if st.button("❌", key=f"inv_item_remove_{i}"):
                    st.session_state.inv_items.pop(i)
                    st.rerun()
        
        # Add new item button
        col1, col2, col3 = st.columns([2,2,2])
        with col1:
            if st.button("➕ Add Item", use_container_width=True):
                st.session_state.inv_items.append({
                    'description': '',
                    'quantity': 1,
                    'unit_price': 0.0,
                    'amount': 0.0
                })
                st.rerun()
        
        # Calculate withholding tax
        platform_fee = 0
        for item in st.session_state.inv_items:
            if 'Platform Fee' in item.get('description', ''):
                platform_fee = item.get('unit_price', 0) * item.get('quantity', 1)
                break
        
        withholding_tax = calculate_withholding_tax(platform_fee, currency)
        
        # Show totals
        st.markdown("---")
        st.markdown("### 📊 Invoice Summary")
        subtotal = sum(item.get('unit_price', 0) * item.get('quantity', 1) for item in st.session_state.inv_items)
        total = subtotal - withholding_tax
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Subtotal", f"{currency} {subtotal:,.2f}")
        with col2:
            if withholding_tax > 0:
                st.metric("Withholding Tax (5%)", f"{currency} {withholding_tax:,.2f}")
            else:
                st.metric("Withholding Tax", "N/A")
        with col3:
            st.metric("Total", f"{currency} {total:,.2f}", delta=f"{currency} {total:,.2f}")
        
        # Save invoice
        if st.button("💾 Save Invoice", type="primary", use_container_width=True):
            if not merchant_name:
                st.error("Please enter a merchant name.")
            elif not period:
                st.error("Please enter the invoice period.")
            elif not st.session_state.inv_items:
                st.error("Please add at least one invoice item.")
            else:
                invoice_id = generate_invoice_id(merchant_name, period)
                
                invoice_data = {
                    'invoice_id': invoice_id,
                    'merchant_name': merchant_name,
                    'merchant_email': merchant_email,
                    'invoice_date': invoice_date.strftime('%d/%m/%Y'),
                    'due_date': due_date.strftime('%d/%m/%Y'),
                    'period': period,
                    'currency': currency,
                    'platform_fee': platform_fee,
                    'kyc_fee': 0,
                    'inbound_fee': 0,
                    'swift_charge': 0,
                    'withholding_tax': withholding_tax,
                    'total_amount': total,
                    'status': status,
                    'payment_details': '''For RTGS;
                    Account Name: Choice Microfinance Bank Limited
                    BANK: CHOICE MICROFINANCE BANK LIMITED
                    BANK CODE: 082
                    ACCOUNT NO: 1000738901
                    BRANCH: HEAD OFFICE
                    BRANCH CODE: 001
                    SWIFT CODE: CHFIKENX
                    
                    Or
                    
                    For TT Transfer/EFT;
                    ACCOUNT NAME: Choice Microfinance Bank Limited
                    BANK: I&M Bank Kenya Ltd
                    BANK CODE: 57
                    ACCOUNT NO: 01004389381250
                    BRANCH: Parklands
                    BRANCH CODE: 010
                    SWIFT CODE: IMBLKENA''',
                    'notes': f'BaaS Service Fee for {period}'
                }
                
                # Update items with final amounts
                items_data = []
                for item in st.session_state.inv_items:
                    items_data.append({
                        'description': item.get('description', ''),
                        'quantity': item.get('quantity', 1),
                        'unit_price': item.get('unit_price', 0),
                        'amount': item.get('unit_price', 0) * item.get('quantity', 1),
                        'currency': currency
                    })
                
                invoice_db.save_invoice(invoice_data, items_data)
                invoice_db.log_history(invoice_id, 'created', 'user', 'Invoice created manually')
                
                st.success(f"✅ Invoice {invoice_id} saved successfully!")
                st.balloons()
                
                # Reset form
                st.session_state.inv_items = []
                st.rerun()
    
    elif module_option == "📊 Analytics":
        st.markdown("### 📊 Invoice Analytics")
        
        # Load data
        df = invoice_db.get_invoices()
        
        if df.empty:
            st.info("No invoice data available. Import or create invoices first.")
        else:
            # Summary metrics
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Total Invoices", len(df))
            with col2:
                sent_df = df[df['status'] == 'Sent']
                st.metric("Sent", len(sent_df))
            with col3:
                paid_df = df[df['status'] == 'Paid']
                st.metric("Paid", len(paid_df))
            with col4:
                total_amount = df['total_amount'].sum()
                st.metric("Total Invoice Amount", f"{total_amount:,.2f}")
            
            # Charts
            col1, col2 = st.columns(2)
            with col1:
                # Status distribution
                status_counts = df['status'].value_counts().reset_index()
                status_counts.columns = ['Status', 'Count']
                fig = px.pie(status_counts, values='Count', names='Status', title='Invoice Status Distribution')
                st.plotly_chart(fig, use_container_width=True)
            
            with col2:
                # Currency distribution
                currency_counts = df['currency'].value_counts().reset_index()
                currency_counts.columns = ['Currency', 'Count']
                fig = px.bar(currency_counts, x='Currency', y='Count', title='Invoices by Currency')
                st.plotly_chart(fig, use_container_width=True)
            
            # Monthly trends
            if 'invoice_date' in df.columns:
                df['month'] = pd.to_datetime(df['invoice_date'], format='%d/%m/%Y', errors='coerce').dt.to_period('M')
                monthly_data = df.groupby('month').agg({
                    'invoice_id': 'count',
                    'total_amount': 'sum'
                }).reset_index()
                monthly_data.columns = ['Month', 'Count', 'Amount']
                
                fig = px.line(monthly_data, x='Month', y=['Count', 'Amount'], title='Monthly Invoice Trends')
                st.plotly_chart(fig, use_container_width=True)
            
            # Top merchants
            st.markdown("### 🏆 Top Merchants by Invoice Amount")
            top_merchants = df.groupby('merchant_name')['total_amount'].sum().sort_values(ascending=False).head(10).reset_index()
            top_merchants.columns = ['Merchant', 'Total Amount']
            st.dataframe(top_merchants, use_container_width=True)
    
    return

# For backward compatibility and direct import
def baas_invoice_management_page():
    baas_invoice_management_app()

if __name__ == "__main__":
    baas_invoice_management_app()