# pages/baas_invoice_generation_page.py
"""
BaaS Invoice Generation Module
===============================
Reads client billing instructions straight from the 'BAAS INCOME' column of the
BaaS Income Final tracker (e.g. "HONGKONG FORTUNETECH BAAS SERVICE FEE FOR JULY 2026
(PLATFORM FEE - 1500usd*129.29; Inbound Fee - 10.56 usd*129.29; Inbound Swift Charge -
7usd*129.29)"), lets the reviewer confirm/adjust the parsed fee breakdown, then
generates, previews, downloads and emails a Choice Bank BaaS invoice (.docx) that
matches the approved invoice template exactly (same letterhead, fonts and layout as
the Ahadi Wireless / Savatech / HongKong FortuneTech invoices).

Design mirrors interfund_bank_reconciliation_page.py: a CUSTOM_CSS block, a small
SQLite-backed persistence layer, session-state driven tabs, and editable
st.data_editor review grids so nothing is ever auto-sent without a human check.
"""

import os
import re
import io
import json
import uuid
import ssl
import base64
import smtplib
import sqlite3
import logging
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from datetime import datetime, timedelta

import streamlit as st
import pandas as pd
import plotly.express as px
import openpyxl

try:
    from fuzzywuzzy import process as fuzzy_process
    from fuzzywuzzy import fuzz as fuzzy_fuzz
except Exception:  # pragma: no cover - fuzzywuzzy is already a dependency elsewhere in the app
    fuzzy_process = None
    fuzzy_fuzz = None

from docx import Document
from docx.shared import Pt, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    from weasyprint import HTML as _WeasyHTML
except Exception:  # pragma: no cover - optional dependency, PDF export degrades gracefully
    _WeasyHTML = None

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ============================================================================================
# Constants & paths
# ============================================================================================
DATA_DIR = "data/uploads"
CACHE_DIR = "data/cache"
ASSETS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets")
BAAS_DB_PATH = "data/baas_invoicing.db"
os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(CACHE_DIR, exist_ok=True)
print(f'Asset path : ', ASSETS_DIR)
LOGO_PATH = os.path.join(ASSETS_DIR, "choicebank_invoice_header.png")
FOOTER_BANNER_PATH = os.path.join(ASSETS_DIR, "choicebank_invoice_footer.png")

WHT_RATE = 0.05            # 5% withholding tax - KES invoices only

# Base font size used throughout the generated .docx invoice (points). Kept as a single
# constant so the whole document scales consistently if it ever needs to change again.
INVOICE_FONT_SIZE_PT = 10


def compute_wht_gross_up(platform_fee_net):
    """
    Grosses up a KES platform fee and computes the withholding tax on it, using the
    exact formula supplied for this workbook (platform_fee is the net figure taken
    straight from the BAAS INCOME text, e.g. the "200,000" in "PLATFORM FEE - 200,000"):

        Withholding Tax     = platform_fee * 5 / 95
        Gross Platform Fee  = platform_fee * 100 / 95
        Net Payable         = platform_fee                    (i.e. 95% of the gross)

    Returns (gross, withholding_tax), both rounded to 2 decimal places.
    """
    gross = round(platform_fee_net * 100 / 95, 2)
    wht = round(platform_fee_net * 5 / 95, 2)
    return gross, wht
DEFAULT_FX_RATE = 129.29   # fallback USD -> KES rate if none can be found in the row text
CURRENCY_SYMBOLS = {"KES": "KES", "USD": "USD", "EUR": "EUR", "GBP": "GBP"}

MONTH_MAP = {
    'JAN': 1, 'JANUARY': 1, 'FEB': 2, 'FEBRUARY': 2, 'MAR': 3, 'MARCH': 3,
    'APR': 4, 'APRIL': 4, 'MAY': 5, 'JUN': 6, 'JUNE': 6, 'JUL': 7, 'JULY': 7,
    'AUG': 8, 'AUGUST': 8, 'SEP': 9, 'SEPT': 9, 'SEPTEMBER': 9, 'OCT': 10,
    'OCTOBER': 10, 'NOV': 11, 'NOVEMBER': 11, 'DEC': 12, 'DECEMBER': 12
}

# Ordered longest-first so "BAAS SERVICE FEE" is matched before the generic "FEE"
FEE_KEYWORDS = [
    'BAAS SERVICE FEE', 'SERVICE FEE FOR', 'PLATFORM FEE FOR',
    'SERVICE FEE', 'PLATFORM FEE', 'FEE FOR'
]

PLATFORM_LABEL_HINTS = ('PLATFORM',)

# Custom CSS - same visual language (gradient header / stat cards / rounded tabs)
# as interfund_bank_reconciliation_page.py so the module feels native to the app.
CUSTOM_CSS = """
<style>
    .baas-main-header {
        background: linear-gradient(135deg, #4B2D8F 0%, #6B4DB5 100%);
        padding: 1.5rem;
        border-radius: 10px;
        margin-bottom: 1.5rem;
        color: white;
    }
    .baas-stat-card {
        background: white;
        padding: 1rem;
        border-radius: 10px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        margin: 0.5rem 0;
        border-left: 4px solid #4B2D8F;
    }
    .baas-warning-card {
        background-color: #fff3cd;
        border-left: 4px solid #f5c842;
        padding: 0.75rem 1rem;
        border-radius: 8px;
        margin: 0.5rem 0;
        font-size: 0.9rem;
    }
    .baas-success-card {
        background-color: #d4edda;
        border-left: 4px solid #28a745;
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
    }
    .baas-invoice-preview {
        background: white;
        border: 1px solid #E2D9F3;
        border-radius: 10px;
        padding: 1.5rem;
        box-shadow: 0 2px 8px rgba(0,0,0,0.08);
        font-family: 'Times New Roman', serif;
    }
    .baas-invoice-preview table { width: 100%; border-collapse: collapse; margin-top: 0.75rem; }
    .baas-invoice-preview th, .baas-invoice-preview td {
        border: 1px solid #333; padding: 6px 10px; text-align: left; font-size: 0.95rem;
    }
    .baas-invoice-preview th { background: #F0EBF9; }
    .stTabs [data-baseweb="tab-list"] {
        gap: 1rem; background-color: #f8f9fa; padding: 0.5rem; border-radius: 10px;
    }
    .stTabs [data-baseweb="tab"] { border-radius: 8px; padding: 0.5rem 1rem; font-weight: 500; }
</style>
"""

# ============================================================================================
# Seed client directory
# Fully known clients (exact banking data lifted from the approved invoice templates) plus the
# remaining active merchants pulled from the 'Merchant_Master (Static)' sheet of the BaaS Income
# workbook (currency + invoice day-of-month only - banking/contact details left blank for the
# reviewer to complete once, after which they are remembered for every future month).
# ============================================================================================
def _blank_client(legal_name, currency, invoice_group, match_keywords):
    return {
        "id": str(uuid.uuid4()),
        "match_keywords": match_keywords,
        "legal_name": legal_name,
        "invoice_prefix": re.sub(r'[^A-Z0-9]', '', legal_name.upper().split()[0]) if legal_name else "",
        "contact_email": "",
        "currency": currency,
        "invoice_group": invoice_group or "5th",
        "rtgs_bank_name": "CHOICE MICROFINANCE BANK LIMITED",
        "rtgs_account_no": "",
        "rtgs_bank_code": "082",
        "rtgs_branch": "HEAD OFFICE",
        "rtgs_branch_code": "001",
        "rtgs_swift": "CHFIKENX",
        "paybill": "",
        "alt_method_label": "For SWIFT Transfer;",
        "alt_bank_name": "I&M Bank Kenya Ltd",
        "alt_account_no": "",
        "alt_bank_code": "57",
        "alt_branch": "Parklands",
        "alt_branch_code": "010",
        "alt_swift": "IMBLKENA",
        "active": "Yes",
    }

_MERCHANT_MASTER_SEED = [
    ("KOASAVE AFRICA LIMITED", "KES", "5th"),
    ("SAVATECH LIMITED", "USD", "5th"),
    ("AHADI WIRELESS LIMITED/Konnect", "KES", "5th"),
    ("COUPON SASA LIMITED", "KES", "5th"),
    ("PAYCLOUD AFRICA SYSTEMS LIMITED", "KES", "5th"),
    ("NIOBI KENYA LIMITED", "KES", "5th"),
    ("SWITCHLINK AFRICA LIMITED", "KES", "5th"),
    ("BOYA LIMITED", "KES", "5th"),
    ("AVENEWS KE LIMITED", "KES", "5th"),
    ("TRIPITACA LIMITED", "KES", "5th"),
    ("RWK& ASSOCIATES CPA(K) LTD", "KES", "5th"),
    ("LIVING TWICE, INC.", "USD", "5th"),
    ("KEYMAN OAK LIMITED", "KES", "5th"),
    ("SENTZEUS LIMITED", "KES", "5th"),
    ("COWDI TECH LIMITED", "KES", "5th"),
    ("WORKPAY AFRICA LIMITED", "KES", "5th"),
    ("KESH LABS LTD", "KES", "5th"),
    ("LUKA MIGRANT TECHNOLOGY LIMITED", "KES", "5th"),
    ("CLOUD NINE PAYMENTS LIMITED", "USD", "5th"),
    ("NUMIDA TECHNOLOGIES KENYA LIMITED", "USD", "5th"),
    ("DUPLO TECHNOLOGIES LIMITED", "USD", "5th"),
    ("QUICKNODE TECHNOLOGIES LIMITED", "KES", "10th"),
    ("FINCRA TECHNOLOGIES LIMITED", "USD", "10th"),
    ("VERTOFX LTD", "USD", "10th"),
    ("ONELOOP (PTE.) LIMITED", "KES", "10th"),
    ("PING PONG GLOBAL HOLDINGS LIMITED", "USD", "10th"),
    ("HONGKONG FORTUNETECH LIMITED", "USD", "10th"),
    ("TEMBOPLUS COMPANY LIMITED", "USD", "10th"),
    ("KUWAPAY MONEY TRANSFER LIMITED", "USD", "10th"),
    ("Barnes Bridge Capital", "KES", "10th"),
    ("CAPI MONEY CANADA LTD", "KES", "10th"),
    ("CRISSCROSS FX LIMITED", "KES", "10th"),
    ("YELLOW PAY LIMITED", "KES", "10th"),
    ("PESASWAP EAST AFRICA LIMITED", "KES", "10th"),
    ("NALA PAYMENTS LIMITED", "KES", "5th"),
]

# Keyword fragments used to fuzzy-match a BAAS INCOME description back to a client record.
_EXTRA_KEYWORDS = {
    "AHADI WIRELESS LIMITED/Konnect": "AHADI,KONNECT",
    "SAVATECH LIMITED": "SAVA,SAVATECH",
    "HONGKONG FORTUNETECH LIMITED": "HONGKONG,FORTUNETECH",
    "RWK& ASSOCIATES CPA(K) LTD": "RWK,RWK AFRICA",
    "KUWAPAY MONEY TRANSFER LIMITED": "KUWAPAY,FESAKI",
    "PING PONG GLOBAL HOLDINGS LIMITED": "PINGPONG,PING PONG",
    "ONELOOP (PTE.) LIMITED": "ONELOOP",
    "VERTOFX LTD": "VERTO",
}


def _default_keywords(legal_name):
    # Generic fintech/banking words are deliberately excluded so they never masquerade as a
    # match for an unrelated client (e.g. "TECH" inside "COWDI TECH LIMITED" must not match
    # "HONGKONG FORTUNETECH LIMITED" just because both strings contain "TECH").
    stop = {"LIMITED", "LTD", "LLC", "INC", "AFRICA", "KENYA", "COMPANY", "HOLDINGS",
            "TECHNOLOGIES", "TECHNOLOGY", "SOLUTIONS", "SYSTEMS", "GLOBAL", "AND", "ASSOCIATES",
            "TECH", "PAY", "PAYMENTS", "PAYMENT", "LABS", "FINANCE", "FINANCIAL", "CAPITAL",
            "MONEY", "BRIDGE", "GROUP", "CORP", "CORPORATION", "SERVICES", "EAST", "PTE",
            "CANADA", "ASIA", "USA", "WORLD", "INTERNATIONAL", "ENTERPRISES", "STUDIOS"}
    tokens = re.split(r'[^A-Za-z]+', legal_name.upper())
    keep = [t for t in tokens if t and t not in stop and len(t) >= 4]
    return ",".join(keep[:3]) if keep else legal_name.upper()


def build_seed_clients():
    seed = []
    for legal_name, currency, group in _MERCHANT_MASTER_SEED:
        row = _blank_client(legal_name, currency, group, _EXTRA_KEYWORDS.get(legal_name, _default_keywords(legal_name)))
        seed.append(row)

    overrides = {
        "AHADI WIRELESS LIMITED/Konnect": dict(
            invoice_prefix="AHADI", contact_email="ben@ahadicorp.com",
            rtgs_account_no="1000238348", alt_method_label="For SWIFT Transfer/Pesalink/Mpesa;",
            alt_account_no="01004389386250", paybill="542542",
        ),
        "SAVATECH LIMITED": dict(
            invoice_prefix="SAVA", contact_email="yoel@sava.africa",
            rtgs_account_no="1000738901", alt_method_label="For SWIFT Transfer;",
            alt_account_no="01004389381250",
        ),
        "HONGKONG FORTUNETECH LIMITED": dict(
            invoice_prefix="HONGKONG", contact_email="joe.yu@1yunhui.com",
            rtgs_account_no="1000738901", alt_method_label="For TT Transfer/EFT;",
            alt_account_no="01004389381250",
        ),
    }
    for row in seed:
        if row["legal_name"] in overrides:
            row.update(overrides[row["legal_name"]])
    return seed


CLIENT_COLUMNS = list(_blank_client("", "KES", "5th", "").keys())


# ============================================================================================
# SQLite persistence
# ============================================================================================
class BaaSInvoiceDB:
    """Lightweight persistence for the client directory and the sent/generated invoice log."""

    def __init__(self, db_path=BAAS_DB_PATH):
        self.db_path = db_path
        self._init_database()

    def _init_database(self):
        conn = sqlite3.connect(self.db_path)
        cur = conn.cursor()
        cols_sql = ", ".join([f'"{c}" TEXT' for c in CLIENT_COLUMNS if c != "id"])
        cur.execute(f'''
            CREATE TABLE IF NOT EXISTS baas_clients (
                id TEXT PRIMARY KEY,
                {cols_sql}
            )
        ''')
        cur.execute('''
            CREATE TABLE IF NOT EXISTS baas_invoice_log (
                id TEXT PRIMARY KEY,
                created_at TEXT,
                invoice_no TEXT,
                client_legal_name TEXT,
                period TEXT,
                currency TEXT,
                subtotal_gross REAL,
                withholding_tax REAL,
                net_total REAL,
                status TEXT,
                sent_to TEXT,
                sent_at TEXT,
                raw_source_text TEXT,
                generated_by TEXT
            )
        ''')
        conn.commit()
        conn.close()

        # Seed the client directory once, on first ever run.
        existing = self.load_clients()
        if existing.empty:
            self.save_clients(pd.DataFrame(build_seed_clients()))

    def load_clients(self):
        conn = sqlite3.connect(self.db_path)
        try:
            df = pd.read_sql_query("SELECT * FROM baas_clients", conn)
        except Exception:
            df = pd.DataFrame(columns=CLIENT_COLUMNS)
        conn.close()
        return df

    def save_clients(self, df):
        if df is None:
            return
        df = df.copy()
        for c in CLIENT_COLUMNS:
            if c not in df.columns:
                df[c] = ""
        df["id"] = df["id"].apply(lambda x: x if isinstance(x, str) and x else str(uuid.uuid4()))
        df = df[CLIENT_COLUMNS].fillna("")
        conn = sqlite3.connect(self.db_path)
        df.to_sql("baas_clients", conn, if_exists="replace", index=False)
        conn.commit()
        conn.close()

    def log_invoice(self, record: dict):
        conn = sqlite3.connect(self.db_path)
        cur = conn.cursor()
        record = dict(record)
        record.setdefault("id", str(uuid.uuid4()))
        record.setdefault("created_at", datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        cols = ["id", "created_at", "invoice_no", "client_legal_name", "period", "currency",
                "subtotal_gross", "withholding_tax", "net_total", "status", "sent_to", "sent_at",
                "raw_source_text", "generated_by"]
        cur.execute(f'''
            INSERT OR REPLACE INTO baas_invoice_log ({",".join(cols)})
            VALUES ({",".join(["?"] * len(cols))})
        ''', [record.get(c) for c in cols])
        conn.commit()
        conn.close()

    def load_log(self):
        conn = sqlite3.connect(self.db_path)
        try:
            df = pd.read_sql_query("SELECT * FROM baas_invoice_log ORDER BY created_at DESC", conn)
        except Exception:
            df = pd.DataFrame()
        conn.close()
        return df


baas_db = BaaSInvoiceDB()


# ============================================================================================
# Parsing engine - turns one 'BAAS INCOME' free-text cell into structured invoice instructions
# ============================================================================================
LINE_RE = re.compile(
    r'^\s*(?P<label>[A-Za-z /&.]+?)\s*[-:]?\s*(?P<amount>[\d,]+\.?\d*)\s*'
    r'(?P<currency>usd|kes|eur|gbp)?\s*(?:[*x]\s*(?P<rate>[\d.]+))?\s*$',
    re.I
)
LINE_RE_DIV = re.compile(
    r'^\s*(?P<label>[A-Za-z /&.]+?)\s+(?P<amount>[\d,]+\.?\d*)\s*/\s*(?P<rate>[\d.]+)\s*$',
    re.I
)


def split_client_and_rest(text):
    """Return (client_name_guess, remainder_of_string_from_the_fee_keyword_onwards)."""
    upper = text.upper()
    best_pos, best_kw = None, None
    for kw in FEE_KEYWORDS:
        m = re.search(re.escape(kw), upper)
        if m and (best_pos is None or m.start() < best_pos):
            best_pos, best_kw = m.start(), kw
    if best_pos is None:
        return text.strip(), text
    client = text[:best_pos].strip(' /\t-')
    return client, text[best_pos:]


def parse_period(rest):
    """Extract (month1, month2_or_None, year_or_None) from the 'FOR <MONTH> <YEAR>' segment."""
    m = re.search(r'FOR\s+([A-Za-z]+)(?:\s*&\s*([A-Za-z]+))?\s*(\d{4})?', rest, re.I)
    if not m:
        m2 = re.search(r'FOR([A-Za-z]+)', rest, re.I)  # tolerate "FORMAY" typos
        if m2:
            return m2.group(1).upper(), None, None
        return None, None, None
    month1 = m.group(1).upper()
    month2 = m.group(2).upper() if m.group(2) else None
    year = int(m.group(3)) if m.group(3) else None
    return month1, month2, year


def extract_parenthetical(text):
    m = re.search(r'\(([^()]*)\)', text)
    return m.group(1) if m else None


def parse_line_items(paren_text):
    """Split the parenthesised breakdown into individual fee line items."""
    items = []
    if not paren_text:
        return items
    for part in [p.strip() for p in paren_text.split(';') if p.strip()]:
        md = LINE_RE_DIV.match(part)
        if md:
            items.append({
                "label": md.group('label').strip(' -'),
                "raw_amount": float(md.group('amount').replace(',', '')),
                "currency_tag": None,
                "rate": float(md.group('rate')),
                "notation": "div",
                "raw": part,
            })
            continue
        m = LINE_RE.match(part)
        if m:
            items.append({
                "label": m.group('label').strip(' -'),
                "raw_amount": float(m.group('amount').replace(',', '')),
                "currency_tag": m.group('currency').upper() if m.group('currency') else None,
                "rate": float(m.group('rate')) if m.group('rate') else None,
                "notation": "mul",
                "raw": part,
            })
        else:
            items.append({"label": part, "raw_amount": None, "currency_tag": None,
                           "rate": None, "notation": None, "raw": part})
    return items


def parse_baas_income_text(raw_text):
    """Top level parse: raw 'BAAS INCOME' cell text -> structured dict."""
    raw_text = str(raw_text or "").strip()
    client_guess, rest = split_client_and_rest(raw_text)
    month1, month2, year = parse_period(rest)
    paren = extract_parenthetical(rest)
    line_items = parse_line_items(paren)

    # Any explicit currency tag or FX-rate multiplier anywhere -> text suggests a USD invoice.
    detected_currency = "KES"
    for it in line_items:
        if it["currency_tag"] == "USD" or it["rate"] is not None:
            detected_currency = "USD"
            break

    return {
        "raw_text": raw_text,
        "client_guess": client_guess,
        "period_month": month1,
        "period_month2": month2,
        "period_year": year,
        "line_items": line_items,
        "detected_currency": detected_currency,
    }


def extract_baas_income_rows(file_bytes, sheet_name):
    """Read column A ('BAAS INCOME') of the chosen sheet and return the candidate free-text rows."""
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    if sheet_name not in wb.sheetnames:
        return []
    ws = wb[sheet_name]
    rows = []
    for r in range(2, ws.max_row + 1):  # row 1 is the header / scratch-calc row
        val = ws.cell(row=r, column=1).value
        if not val or not isinstance(val, str):
            continue
        text = val.strip()
        if len(text) < 12:
            continue
        if "FEE" not in text.upper():
            continue
        rows.append({"row": r, "text": text})
    return rows


def guess_gl_sheet_names(file_bytes):
    """List sheets that look like a monthly 'GL Breakdown <Month> <Year>' tab, most-recent first."""
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True)
    candidates = [s for s in wb.sheetnames if "GL Breakdown" in s or "BAAS Finance Tracker" in s]
    return candidates or wb.sheetnames


FUZZY_MATCH_THRESHOLD = 90  # deliberately high: a wrong match silently attaches someone
                            # else's bank account/currency to an invoice, so "no match" (which
                            # the reviewer must resolve by hand) is always the safer failure mode.


def match_client(client_guess, clients_df):
    """Match a parsed client name guess against the client directory. Returns (row_or_None, score)."""
    if clients_df is None or clients_df.empty or not client_guess:
        return None, 0
    choices = {}
    for _, row in clients_df.iterrows():
        keys = [row.get("legal_name", "")] + str(row.get("match_keywords", "")).split(",")
        for k in keys:
            k = k.strip()
            if k:
                choices[k] = row

    guess_upper = client_guess.upper()
    # Fast path: direct substring match on a keyword. When several keywords match, the longest
    # (most specific) one wins so a generic fragment ("TECH") can never beat a specific one
    # ("FORTUNETECH") just because it happened to be inserted first.
    best_k, best_row = None, None
    for k, row in choices.items():
        ku = k.upper()
        if ku and (ku in guess_upper or guess_upper in ku):
            if best_k is None or len(ku) > len(best_k):
                best_k, best_row = ku, row
    if best_row is not None:
        return best_row, 100

    # Fuzzy fallback is deliberately restricted to full legal names (never the short keyword
    # fragments - a 4-letter fragment like "SAVA" can score misleadingly high against a
    # completely unrelated longer name) and uses a strict threshold + scorer.
    if fuzzy_process is None:
        return None, 0
    legal_names = clients_df["legal_name"].dropna().tolist()
    if not legal_names:
        return None, 0
    scorer = fuzzy_fuzz.token_sort_ratio if fuzzy_fuzz else None
    best = fuzzy_process.extractOne(client_guess, legal_names, scorer=scorer) if scorer \
        else fuzzy_process.extractOne(client_guess, legal_names)
    if best and best[1] >= FUZZY_MATCH_THRESHOLD:
        match = clients_df[clients_df["legal_name"] == best[0]].iloc[0]
        return match, best[1]
    return None, best[1] if best else 0


def resolve_invoice_amounts(parsed, client_row=None, currency_override=None, fx_rate_override=None):
    """
    Turn parsed line items into final invoice amounts.

    Resolution rules
    -----------------
    * No FX conversion is ever applied to a line item's amount. Whatever number
      appears in the BAAS INCOME text for a given line - however it is written
      (a bare number, 'usd'-tagged, '*rate' or '/rate') - is used exactly as-is
      as that line's amount in the invoice's own currency. Any '*rate' / '/rate'
      suffix is purely the tracker's own internal bookkeeping arithmetic and is
      ignored for invoicing purposes; it is never multiplied or divided into the
      line amount.
    * Invoice currency: an explicit override wins; otherwise the client directory's
      saved currency wins; otherwise it is inferred from the text (any usd tag/rate
      present -> USD); otherwise it defaults to KES, per policy.
    * Withholding tax (KES invoices only) is applied to the line item labelled
      'PLATFORM ...', using the exact formula supplied for this workbook, where
      `platform_fee` is the net amount taken straight from the BAAS INCOME text:
          Withholding Tax   = platform_fee * 5 / 95
          Gross Platform Fee = platform_fee * 100 / 95
          Net Payable (this line) = platform_fee   (i.e. 95% of the gross)
    """
    warnings = []
    line_items = parsed["line_items"]

    text_currency = parsed["detected_currency"]
    master_currency = (client_row.get("currency") if client_row is not None else None) or None
    if currency_override:
        invoice_currency = currency_override
    elif master_currency:
        invoice_currency = master_currency
        if master_currency != text_currency:
            warnings.append(
                f"Client directory says {master_currency}, but the text looks like {text_currency}. "
                f"Using {master_currency} (override the Currency column below if this is wrong)."
            )
    else:
        invoice_currency = text_currency  # already defaults to KES when nothing is detected

    fx_rate = fx_rate_override or next((it["rate"] for it in line_items if it.get("rate")), None) or DEFAULT_FX_RATE

    resolved = []
    for it in line_items:
        if it["raw_amount"] is None:
            warnings.append(f"Could not parse amount from: '{it['raw']}' - please fill in manually.")
            resolved.append({"label": it["label"], "amount": 0.0, "raw": it["raw"]})
            continue

        # No conversion, ever: the amount in the text IS the line amount, regardless of
        # currency tag, notation ('*rate' / '/rate') or the invoice's own currency.
        amount = it["raw_amount"]

        is_platform = any(hint in it["label"].upper() for hint in PLATFORM_LABEL_HINTS)
        resolved.append({"label": it["label"], "amount": round(amount, 2), "raw": it["raw"], "is_platform": is_platform})

    # Sanity flag: a platform-fee line that is wildly larger than the others is almost certainly
    # a leftover KES figure the preparer meant to flag - surface it instead of silently
    # invoicing an unexpected amount (amounts are never auto-converted, so this is the main
    # safety net for that data-entry pattern).
    for r in resolved:
        if r.get("is_platform") and invoice_currency == "USD" and r["amount"] > 20000:
            warnings.append(
                f"Platform fee parsed as {r['amount']:,.2f} USD - unusually large. "
                f"Double-check this against the source text ('{r['raw']}') before sending."
            )

    subtotal_gross = 0.0
    total_wht = 0.0
    final_items = []
    for r in resolved:
        gross = r["amount"]
        wht = 0.0
        if r.get("is_platform") and invoice_currency == "KES":
            gross, wht = compute_wht_gross_up(r["amount"])
        subtotal_gross += gross
        total_wht += wht
        final_items.append({"label": r["label"], "amount": r["amount"], "gross": gross, "wht": wht,
                             "is_platform": r.get("is_platform", False)})

    net_total = round(subtotal_gross - total_wht, 2)
    return {
        "invoice_currency": invoice_currency,
        "fx_rate_used": fx_rate,
        "line_items": final_items,
        "subtotal_gross": round(subtotal_gross, 2),
        "total_wht": round(total_wht, 2),
        "net_total": net_total,
        "warnings": warnings,
    }


def compute_invoice_dates(period_month, period_year, invoice_group):
    """Invoice is raised on the 5th/10th of the service month itself; due 7 days later."""
    year = period_year or datetime.now().year
    month = MONTH_MAP.get((period_month or "").upper())
    if not month:
        month, year = datetime.now().month, datetime.now().year
    day = 10 if str(invoice_group).strip().startswith("10") else 5
    try:
        invoice_date = datetime(year, month, day)
    except ValueError:
        invoice_date = datetime(year, month, 1)
    due_date = invoice_date + timedelta(days=7)
    return invoice_date, due_date


def build_invoice_no(invoice_prefix, invoice_date):
    return f"{invoice_prefix}/{invoice_date.strftime('%Y%m')}"


def format_money(amount, currency):
    return f"{currency} {amount:,.2f}"


def format_label(label):
    """Title-case a parsed label, but keep short all-caps acronyms (KYB, KYC) intact."""
    words = []
    for word in label.strip().split():
        clean = re.sub(r'[^A-Za-z]', '', word)
        if clean.isupper() and 2 <= len(clean) <= 4:
            words.append(word)
        else:
            words.append(word.title())
    return " ".join(words)


# ============================================================================================
# DOCX generation - rebuilds the approved invoice template exactly (letterhead, fonts, layout)
# ============================================================================================
def _add_run(paragraph, text, bold=False, italic=False, size=INVOICE_FONT_SIZE_PT, font_name="Times New Roman"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    return run


def _tab_paragraph(document, segments, style=None, tab_stop_inches=None):
    """segments: list of (text, bold) tuples, tabs already embedded in the text."""
    p = document.add_paragraph(style=style)
    if tab_stop_inches:
        p.paragraph_format.tab_stops.add_tab_stop(Emu(int(tab_stop_inches * 914400)))
    for text, bold in segments:
        _add_run(p, text, bold=bold)
    return p


def _picture_header_footer(document, image_path, container):
    if not image_path or not os.path.exists(image_path):
        return
    p = container.paragraphs[0] if container.paragraphs else container.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run()
    usable_width = document.sections[0].page_width - document.sections[0].left_margin - document.sections[0].right_margin
    run.add_picture(image_path, width=usable_width)


def generate_invoice_docx(client, invoice_meta, computed):
    """
    client: dict-like row from the client directory
    invoice_meta: {invoice_no, invoice_date, due_date, period_label}
    computed: output of resolve_invoice_amounts()
    Returns: BytesIO of the finished .docx
    """
    document = Document()
    # python-docx's built-in template defaults every paragraph to ~10pt space-after and 1.15x
    # line spacing (the modern Word default). The approved invoice template has neither, so
    # left alone this alone is enough extra height to push the whole letter onto a second page.
    normal_style = document.styles['Normal']
    normal_style.font.size = Pt(INVOICE_FONT_SIZE_PT)
    normal_style.font.name = "Times New Roman"
    normal_style.paragraph_format.space_after = Pt(0)
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.line_spacing = 1.0
    section = document.sections[0]
    section.page_width = Emu(7543800)
    section.page_height = Emu(10687050)
    section.left_margin = Emu(495300)
    section.right_margin = Emu(1066800)
    section.top_margin = Emu(495300)
    section.bottom_margin = Emu(0)
    section.header_distance = Emu(539750)
    section.footer_distance = Emu(457200)

    _picture_header_footer(document, LOGO_PATH, section.header)
    _picture_header_footer(document, FOOTER_BANNER_PATH, section.footer)

    currency = computed["invoice_currency"]

    # A fixed tab stop (rather than a run of literal tab characters) keeps "Invoice No:" /
    # "Date:" aligned at a consistent horizontal position no matter how long the client's
    # legal name or email address is, so a long name can never push the second column onto
    # a wrapped second line the way repeated \t characters would.
    HEADER_TAB_STOP_IN = 3.7
    document.add_paragraph()
    _tab_paragraph(document, [
        (f"To: {client.get('legal_name', '').upper()}\t", False),
        (f"Invoice No: {invoice_meta['invoice_no']}", False),
    ], tab_stop_inches=HEADER_TAB_STOP_IN)
    document.add_paragraph()
    _tab_paragraph(document, [
        (f"Email: {client.get('contact_email', '')}\t", False),
        (f"Date: {invoice_meta['invoice_date'].strftime('%d/%m/%Y')}", False),
    ], tab_stop_inches=HEADER_TAB_STOP_IN)
    _tab_paragraph(document, [
        ("\t", False),
        (f"Due Date: {invoice_meta['due_date'].strftime('%d/%m/%Y')}", False),
    ], tab_stop_inches=HEADER_TAB_STOP_IN)
    for _ in range(5):
        document.add_paragraph()

    # Description / Amount table
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.columns[0].width = Emu(4029075)
    table.columns[1].width = Emu(2089785)

    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["DESCRIPTION", "AMOUNT"]):
        cell.text = ""
        p = cell.paragraphs[0]
        _add_run(p, text, bold=True)

    body_row = table.add_row().cells
    body_row[0].text = ""
    body_row[1].text = ""
    desc_p = body_row[0].paragraphs[0]
    amt_p = body_row[1].paragraphs[0]

    period_label = invoice_meta["period_label"]
    first = True
    for item in computed["line_items"]:
        label = format_label(item["label"])
        if item.get("is_platform"):
            line_desc = f"Monthly Platform Fee ({period_label})"
        else:
            line_desc = label
        if not first:
            desc_p = body_row[0].add_paragraph()
            amt_p = body_row[1].add_paragraph()
        _add_run(desc_p, line_desc)
        _add_run(amt_p, format_money(item["gross"], currency))
        first = False
        if item.get("is_platform") and item["wht"] > 0:
            wht_desc_p = body_row[0].add_paragraph()
            wht_amt_p = body_row[1].add_paragraph()
            _add_run(wht_desc_p, f"Less: Withholding Tax ({int(WHT_RATE * 100)}%) ")
            _add_run(wht_amt_p, format_money(item["wht"], currency))

    total_row = table.add_row().cells
    total_row[0].text = ""
    total_row[1].text = ""
    _add_run(total_row[0].paragraphs[0], "TOTAL NET PAYABLE", bold=True)
    _add_run(total_row[1].paragraphs[0], format_money(computed["net_total"], currency))

    document.add_paragraph()
    p = document.add_paragraph()
    _add_run(p, "Payment", bold=True)
    _add_run(p, " Details: ", bold=True)
    document.add_paragraph()

    p = document.add_paragraph()
    _add_run(p, "For RTGS;", italic=True)
    _tab_paragraph(document, [("Account Name:\t", True), ("Choice Microfinance Bank Limited", False)])
    _tab_paragraph(document, [("BANK:\t\t", True), (client.get("rtgs_bank_name", "CHOICE MICROFINANCE BANK LIMITED") + " ", False)])
    _tab_paragraph(document, [("BANK CODE:\t", True), (str(client.get("rtgs_bank_code", "082")) + " ", False)])
    _tab_paragraph(document, [("ACCOUNT NO:\t", True), (str(client.get("rtgs_account_no", "")) + " ", False)])
    _tab_paragraph(document, [("BRANCH:\t\t", True), (client.get("rtgs_branch", "HEAD OFFICE") + " ", False)])
    _tab_paragraph(document, [("BRANCH CODE:\t", True), (str(client.get("rtgs_branch_code", "001")) + " ", False)])
    _tab_paragraph(document, [("SWIFT CODE:\t", True), (client.get("rtgs_swift", "CHFIKENX") + " ", False)])
    narration = f"{client.get('invoice_prefix', '')} Monthly BaaS Service Fee {invoice_meta['invoice_date'].strftime('%Y%m')}"
    _tab_paragraph(document, [("NARRATION:\t\t", True), (narration, False)])
    document.add_paragraph()
    p = document.add_paragraph()
    _add_run(p, "Or")
    document.add_paragraph()

    p = document.add_paragraph()
    _add_run(p, client.get("alt_method_label", "For SWIFT Transfer;"), italic=True)
    _tab_paragraph(document, [("ACCOUNT NAME:\t", True), ("Choice Microfinance Bank Limited", False)])
    _tab_paragraph(document, [("BANK: \t\t", True), (client.get("alt_bank_name", "I&M Bank Kenya Ltd"), False)])
    _tab_paragraph(document, [("BANK CODE: \t", True), (str(client.get("alt_bank_code", "57")), False)])
    _tab_paragraph(document, [("ACCOUNT NO: \t", True), (str(client.get("alt_account_no", "")), False)])
    _tab_paragraph(document, [("BRANCH: \t\t", True), (client.get("alt_branch", "Parklands"), False)])
    _tab_paragraph(document, [("BRANCH CODE: \t", True), (str(client.get("alt_branch_code", "010")), False)])
    _tab_paragraph(document, [("SWIFT CODE: \t", True), (client.get("alt_swift", "IMBLKENA"), False)])
    if client.get("paybill"):
        _tab_paragraph(document, [("Paybill:\t\t", True), (str(client.get("paybill")), False)])
    _tab_paragraph(document, [("NARRATION:\t", True), (narration, False)])

    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf


# ============================================================================================
# PDF generation - renders the same invoice as HTML (CSS Paged Media) and converts via
# WeasyPrint. Kept as a distinct, independent code path from the .docx builder above so a
# PDF-export bug can never affect the approved .docx template.
# ============================================================================================
def _image_data_uri(image_path):
    if not image_path or not os.path.exists(image_path):
        return None
    ext = os.path.splitext(image_path)[1].lstrip('.').lower() or "png"
    with open(image_path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode("ascii")
    return f"data:image/{ext};base64,{b64}"


def build_invoice_html(client, invoice_meta, computed):
    """Build a standalone HTML document for the invoice, styled to mirror the .docx layout."""
    currency = computed["invoice_currency"]
    period_label = invoice_meta["period_label"]

    rows_html = ""
    for item in computed["line_items"]:
        label = format_label(item["label"])
        if item.get("is_platform"):
            label = f"Monthly Platform Fee ({period_label})"
        rows_html += f"<tr><td>{label}</td><td class='amt'>{format_money(item['gross'], currency)}</td></tr>\n"
        if item.get("is_platform") and item["wht"] > 0:
            rows_html += (f"<tr><td>Less: Withholding Tax ({int(WHT_RATE * 100)}%)</td>"
                          f"<td class='amt'>{format_money(item['wht'], currency)}</td></tr>\n")

    logo_uri = _image_data_uri(LOGO_PATH)
    footer_uri = _image_data_uri(FOOTER_BANNER_PATH)
    header_img_html = f"<img src='{logo_uri}' class='letterhead-img' />" if logo_uri else ""
    footer_img_html = f"<img src='{footer_uri}' class='letterhead-img' />" if footer_uri else ""

    narration = f"{client.get('invoice_prefix', '')} Monthly BaaS Service Fee {invoice_meta['invoice_date'].strftime('%Y%m')}"
    paybill_html = ""
    if client.get("paybill"):
        paybill_html = f"<div class='pay-line'><span class='pay-label'>Paybill:</span> {client.get('paybill')}</div>"

    html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8" />
<style>
    @page {{
        size: A4;
        margin: 0.7in;
    }}

    body {{
        font-family: 'Times New Roman', Times, serif;
        font-size: 11pt;
        line-height: 1.6;
        color: #1a1a1a;
        margin: 0;
        padding: 0;
    }}

    .letterhead-img {{ width: 100%; display: block; margin-bottom: 12px; }}
    .footer-img {{ width: 100%; display: block; margin-top: 24px; }}

    .invoice-meta p {{ margin: 4px 0; }}
    .invoice-meta .row {{ display: flex; justify-content: space-between; }}

    table.invoice-table {{
        width: 100%;
        border-collapse: collapse;
        margin: 20px 0;
    }}
    table.invoice-table th, table.invoice-table td {{
        border: 1px solid #333;
        padding: 6px 10px;
        text-align: left;
    }}
    table.invoice-table th {{ background: #F0EBF9; }}
    table.invoice-table td.amt {{ text-align: right; }}
    table.invoice-table tr.total td {{ font-weight: bold; }}

    .payment-details .heading {{ font-weight: bold; margin-top: 16px; }}
    .payment-details .method {{ font-style: italic; margin-top: 10px; }}
    .pay-line {{ margin: 2px 0; }}
    .pay-label {{ font-weight: bold; display: inline-block; min-width: 110px; }}
</style>
</head>
<body>
    {header_img_html}

    <div class="invoice-meta">
        <div class="row">
            <p><b>To:</b> {client.get('legal_name', '').upper()}</p>
            <p><b>Invoice No:</b> {invoice_meta['invoice_no']}</p>
        </div>
        <div class="row">
            <p><b>Email:</b> {client.get('contact_email', '')}</p>
            <p><b>Date:</b> {invoice_meta['invoice_date'].strftime('%d/%m/%Y')}</p>
        </div>
        <div class="row">
            <p>&nbsp;</p>
            <p><b>Due Date:</b> {invoice_meta['due_date'].strftime('%d/%m/%Y')}</p>
        </div>
    </div>

    <table class="invoice-table">
        <tr><th>DESCRIPTION</th><th>AMOUNT</th></tr>
        {rows_html}
        <tr class="total"><td>TOTAL NET PAYABLE</td><td class="amt">{format_money(computed['net_total'], currency)}</td></tr>
    </table>

    <div class="payment-details">
        <div class="heading">Payment Details:</div>

        <div class="method">For RTGS;</div>
        <div class="pay-line"><span class="pay-label">Account Name:</span> Choice Microfinance Bank Limited</div>
        <div class="pay-line"><span class="pay-label">Bank:</span> {client.get('rtgs_bank_name', 'CHOICE MICROFINANCE BANK LIMITED')}</div>
        <div class="pay-line"><span class="pay-label">Bank Code:</span> {client.get('rtgs_bank_code', '082')}</div>
        <div class="pay-line"><span class="pay-label">Account No:</span> {client.get('rtgs_account_no', '')}</div>
        <div class="pay-line"><span class="pay-label">Branch:</span> {client.get('rtgs_branch', 'HEAD OFFICE')}</div>
        <div class="pay-line"><span class="pay-label">Branch Code:</span> {client.get('rtgs_branch_code', '001')}</div>
        <div class="pay-line"><span class="pay-label">Swift Code:</span> {client.get('rtgs_swift', 'CHFIKENX')}</div>
        <div class="pay-line"><span class="pay-label">Narration:</span> {narration}</div>

        <p>Or</p>

        <div class="method">{client.get('alt_method_label', 'For SWIFT Transfer;')}</div>
        <div class="pay-line"><span class="pay-label">Account Name:</span> Choice Microfinance Bank Limited</div>
        <div class="pay-line"><span class="pay-label">Bank:</span> {client.get('alt_bank_name', 'I&M Bank Kenya Ltd')}</div>
        <div class="pay-line"><span class="pay-label">Bank Code:</span> {client.get('alt_bank_code', '57')}</div>
        <div class="pay-line"><span class="pay-label">Account No:</span> {client.get('alt_account_no', '')}</div>
        <div class="pay-line"><span class="pay-label">Branch:</span> {client.get('alt_branch', 'Parklands')}</div>
        <div class="pay-line"><span class="pay-label">Branch Code:</span> {client.get('alt_branch_code', '010')}</div>
        <div class="pay-line"><span class="pay-label">Swift Code:</span> {client.get('alt_swift', 'IMBLKENA')}</div>
        {paybill_html}
        <div class="pay-line"><span class="pay-label">Narration:</span> {narration}</div>
    </div>

    {f'<img src="{footer_uri}" class="footer-img" />' if footer_uri else ''}
</body>
</html>"""
    return html


def generate_invoice_pdf(client, invoice_meta, computed):
    """
    Renders the invoice HTML (see build_invoice_html) to PDF via WeasyPrint.
    Returns a BytesIO of the finished .pdf, or None if WeasyPrint isn't installed
    (the caller should fall back to offering the .docx only in that case).
    """
    if _WeasyHTML is None:
        return None
    html_str = build_invoice_html(client, invoice_meta, computed)
    buf = io.BytesIO()
    _WeasyHTML(string=html_str, base_url=ASSETS_DIR).write_pdf(buf)
    buf.seek(0)
    return buf


# ============================================================================================
# Email sending
# ============================================================================================
def send_invoice_email(smtp_cfg, to_email, cc_emails, subject, body, attachment_bytes, attachment_name):
    """smtp_cfg: {'host','port','username','password','use_tls'}"""
    msg = MIMEMultipart()
    msg['From'] = smtp_cfg['username']
    msg['To'] = to_email
    if cc_emails:
        msg['Cc'] = cc_emails
    msg['Subject'] = subject
    msg.attach(MIMEText(body, 'plain'))

    part = MIMEApplication(attachment_bytes, Name=attachment_name)
    part['Content-Disposition'] = f'attachment; filename="{attachment_name}"'
    msg.attach(part)

    recipients = [to_email] + [e.strip() for e in cc_emails.split(',') if e.strip()] if cc_emails else [to_email]

    context = ssl.create_default_context()
    with smtplib.SMTP(smtp_cfg['host'], int(smtp_cfg['port'])) as server:
        if smtp_cfg.get('use_tls', True):
            server.starttls(context=context)
        server.login(smtp_cfg['username'], smtp_cfg['password'])
        server.sendmail(smtp_cfg['username'], recipients, msg.as_string())


# ============================================================================================
# Streamlit UI
# ============================================================================================
def _invoice_preview_html(client, invoice_meta, computed):
    currency = computed["invoice_currency"]
    rows_html = ""
    for item in computed["line_items"]:
        label = format_label(item["label"])
        if item.get("is_platform"):
            label = f"Monthly Platform Fee ({invoice_meta['period_label']})"
        rows_html += f"<tr><td>{label}</td><td>{format_money(item['gross'], currency)}</td></tr>"
        if item.get("is_platform") and item["wht"] > 0:
            rows_html += (f"<tr><td>Less: Withholding Tax ({int(WHT_RATE*100)}%)</td>"
                          f"<td>{format_money(item['wht'], currency)}</td></tr>")

    return f"""
    <div class="baas-invoice-preview">
        <p><b>To:</b> {client.get('legal_name','').upper()} &nbsp;&nbsp;&nbsp;&nbsp;
           <b>Invoice No:</b> {invoice_meta['invoice_no']}</p>
        <p><b>Email:</b> {client.get('contact_email','')} &nbsp;&nbsp;&nbsp;&nbsp;
           <b>Date:</b> {invoice_meta['invoice_date'].strftime('%d/%m/%Y')} &nbsp;&nbsp;&nbsp;&nbsp;
           <b>Due:</b> {invoice_meta['due_date'].strftime('%d/%m/%Y')}</p>
        <table>
            <tr><th>DESCRIPTION</th><th>AMOUNT</th></tr>
            {rows_html}
            <tr><td><b>TOTAL NET PAYABLE</b></td><td><b>{format_money(computed['net_total'], currency)}</b></td></tr>
        </table>
    </div>
    """


def _init_session_state():
    if 'baas_invoicing_initialized' not in st.session_state:
        st.session_state.baas_parsed_rows = pd.DataFrame()
        st.session_state.baas_income_bytes = None
        st.session_state.baas_generated_docs = {}   # row_key -> BytesIO
        st.session_state.baas_smtp_cfg = {"host": "smtp.office365.com", "port": 587, "username": "", "password": "", "use_tls": True}
        st.session_state.baas_invoicing_initialized = True


def _row_key(row):
    return f"{row.get('row')}_{row.get('raw_text', '')[:20]}"


def baas_invoice_generation_app():
    st.markdown(CUSTOM_CSS, unsafe_allow_html=True)
    _init_session_state()

    tab_generate, tab_directory, tab_log, tab_stats = st.tabs(
        ["📤 Generate Invoices", "🗂️ Client Directory", "📨 Sent / Generated Log", "📊 Statistics"]
    )

    # ------------------------------------------------------------------ Generate Invoices ----
    with tab_generate:
        st.markdown("### 📅 Load BaaS Income Source")
        col1, col2 = st.columns([2, 1])
        with col1:
            uploaded = st.file_uploader(
                "Upload the 'BaaS Income Final' workbook", type=["xlsx"], key="baas_income_uploader"
            )
            if uploaded is not None:
                st.session_state.baas_income_bytes = uploaded.read()
        with col2:
            st.caption("Reads the free-text 'BAAS INCOME' column and parses each client's fee instructions.")

        if not st.session_state.baas_income_bytes:
            st.info("Upload the BaaS Income Final workbook to begin.")
        else:
            sheet_options = guess_gl_sheet_names(st.session_state.baas_income_bytes)
            sheet_name = st.selectbox("Select the month's GL Breakdown sheet:", options=sheet_options, key="baas_sheet_select")

            if st.button("🔍 Parse BAAS INCOME Rows", type="primary", key="baas_parse_btn"):
                raw_rows = extract_baas_income_rows(st.session_state.baas_income_bytes, sheet_name)
                clients_df = baas_db.load_clients()
                records = []
                for rr in raw_rows:
                    parsed = parse_baas_income_text(rr["text"])
                    client_row, score = match_client(parsed["client_guess"], clients_df)
                    invoice_group = client_row.get("invoice_group", "5th") if client_row is not None else "5th"
                    invoice_date, due_date = compute_invoice_dates(parsed["period_month"], parsed["period_year"], invoice_group)
                    computed = resolve_invoice_amounts(parsed, client_row)
                    records.append({
                        "row": rr["row"],
                        "select": False,
                        "raw_text": rr["text"],
                        "client_guess": parsed["client_guess"],
                        "matched_client": client_row.get("legal_name") if client_row is not None else "",
                        "match_score": score,
                        "period_month": parsed["period_month"] or "",
                        "period_year": parsed["period_year"] or datetime.now().year,
                        "currency": computed["invoice_currency"],
                        # NET platform fee (pre-WHT, pre-gross-up) - i.e. exactly the figure
                        # taken from the BAAS INCOME text. This is what compute_wht_gross_up()
                        # expects as its input, so it must stay net here: showing the already
                        # grossed-up amount in this column would cause the generation step
                        # below to gross it up a second time.
                        "platform_fee": next((i["amount"] for i in computed["line_items"] if i["is_platform"]), 0.0),
                        "other_fees": round(sum(i["gross"] for i in computed["line_items"] if not i["is_platform"]), 2),
                        "withholding_tax": computed["total_wht"],
                        "net_total": computed["net_total"],
                        "warnings": " | ".join(computed["warnings"]),
                    })
                st.session_state.baas_parsed_rows = pd.DataFrame(records)
                st.session_state.baas_generated_docs = {}

            if not st.session_state.baas_parsed_rows.empty:
                st.markdown("### 🧾 Review Parsed Invoices")
                st.markdown(
                    "<div class='baas-warning-card'>Auto-parsed from free text - always check "
                    "<b>Currency</b>, <b>Platform Fee</b> and <b>Warnings</b> before generating. "
                    "Amounts with no currency tag default to KES; withholding tax (5% of the "
                    "grossed-up platform fee) only applies when the final currency is KES.</div>",
                    unsafe_allow_html=True
                )

                edited = st.data_editor(
                    st.session_state.baas_parsed_rows,
                    use_container_width=True,
                    height=min(480, 60 + 38 * len(st.session_state.baas_parsed_rows)),
                    column_config={
                        "select": st.column_config.CheckboxColumn("Generate?"),
                        "currency": st.column_config.SelectboxColumn("Currency", options=["KES", "USD", "EUR", "GBP"]),
                        "platform_fee": st.column_config.NumberColumn(
                            "Platform Fee (Net, pre-WHT)", format="%.2f",
                            help="The net platform fee as it appears in the source text, before "
                                 "the withholding-tax gross-up. This is the figure the invoice's "
                                 "gross Monthly Platform Fee and 5% WHT line are calculated from."
                        ),
                        "other_fees": st.column_config.NumberColumn("Other Fees", format="%.2f"),
                        "withholding_tax": st.column_config.NumberColumn("WHT (5%)", format="%.2f"),
                        "net_total": st.column_config.NumberColumn("Net Total", format="%.2f"),
                    },
                    disabled=["row", "raw_text", "client_guess", "matched_client", "match_score", "warnings"],
                    key="baas_review_editor",
                )
                st.session_state.baas_parsed_rows = edited

                selected = edited[edited["select"] == True]
                st.markdown(f"**{len(selected)} row(s) selected for generation.**")

                if st.button("🧾 Generate Selected Invoices", type="primary", key="baas_generate_btn") and not selected.empty:
                    clients_df = baas_db.load_clients()
                    for _, row in selected.iterrows():
                        client_row = None
                        if row["matched_client"]:
                            match = clients_df[clients_df["legal_name"] == row["matched_client"]]
                            if not match.empty:
                                client_row = match.iloc[0]
                        if client_row is None:
                            client_row = pd.Series(_blank_client(row["client_guess"] or "UNKNOWN CLIENT", row["currency"], "5th", ""))

                        parsed = parse_baas_income_text(row["raw_text"])
                        computed = resolve_invoice_amounts(parsed, client_row, currency_override=row["currency"])
                        # Let manual edits to platform_fee / other_fees in the grid override the parse.
                        # row["platform_fee"] is the NET figure (see the column's help text above) -
                        # re-parsing it here re-applies compute_wht_gross_up() to that same net
                        # figure, so an unedited row reproduces exactly what was already computed
                        # at parse time instead of grossing it up a second time.
                        if computed["line_items"]:
                            for item in computed["line_items"]:
                                if item["is_platform"] and row.get("platform_fee"):
                                    net_desired = float(row["platform_fee"])
                                    if row["currency"] == "KES":
                                        item["gross"], item["wht"] = compute_wht_gross_up(net_desired)
                                    else:
                                        # Amounts are never converted / grossed-up outside KES -
                                        # the figure in the grid is used exactly as entered.
                                        item["gross"], item["wht"] = net_desired, 0.0
                            computed["subtotal_gross"] = round(sum(i["gross"] for i in computed["line_items"]), 2)
                            computed["total_wht"] = round(sum(i["wht"] for i in computed["line_items"]), 2)
                            computed["net_total"] = round(computed["subtotal_gross"] - computed["total_wht"], 2)

                        invoice_group = client_row.get("invoice_group", "5th")
                        invoice_date, due_date = compute_invoice_dates(row["period_month"], int(row["period_year"]), invoice_group)
                        invoice_no = build_invoice_no(client_row.get("invoice_prefix") or "INV", invoice_date)
                        period_label = f"1st {row['period_month'].title()} – {invoice_date.replace(day=28).strftime('%d')}th {row['period_month'].title()} {invoice_date.year}"
                        invoice_meta = {"invoice_no": invoice_no, "invoice_date": invoice_date,
                                         "due_date": due_date, "period_label": period_label}

                        docx_buf = generate_invoice_docx(client_row, invoice_meta, computed)
                        key = _row_key(row)
                        st.session_state.baas_generated_docs[key] = {
                            "buf": docx_buf, "client": client_row, "invoice_meta": invoice_meta,
                            "computed": computed, "raw_text": row["raw_text"],
                        }
                        baas_db.log_invoice({
                            "invoice_no": invoice_no,
                            "client_legal_name": client_row.get("legal_name", ""),
                            "period": f"{row['period_month']} {row['period_year']}",
                            "currency": computed["invoice_currency"],
                            "subtotal_gross": computed["subtotal_gross"],
                            "withholding_tax": computed["total_wht"],
                            "net_total": computed["net_total"],
                            "status": "Generated",
                            "sent_to": "", "sent_at": "",
                            "raw_source_text": row["raw_text"],
                            "generated_by": st.session_state.get("user", {}).get("username", "unknown_user"),
                        })
                    st.success(f"✅ Generated {len(selected)} invoice(s). Scroll down to preview, download or email each one.")

            if st.session_state.baas_generated_docs:
                st.markdown("### 📄 Generated Invoices")
                for key, doc_info in st.session_state.baas_generated_docs.items():
                    client, invoice_meta, computed = doc_info["client"], doc_info["invoice_meta"], doc_info["computed"]
                    with st.expander(f"🧾 {invoice_meta['invoice_no']} — {client.get('legal_name','')} "
                                      f"({format_money(computed['net_total'], computed['invoice_currency'])})", expanded=False):
                        st.markdown(_invoice_preview_html(client, invoice_meta, computed), unsafe_allow_html=True)
                        if computed["warnings"]:
                            for w in computed["warnings"]:
                                st.warning(w)

                        col_a, col_b, col_c, col_d = st.columns([1, 1, 1, 1])
                        with col_a:
                            st.download_button(
                                "⬇️ Download .docx", data=doc_info["buf"].getvalue(),
                                file_name=f"{client.get('legal_name','Invoice').replace(' ', '_')}_{invoice_meta['invoice_no'].replace('/', '_')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"dl_{key}",
                            )
                        with col_b:
                            # PDF is rendered on demand (rather than at generation time) so the
                            # WeasyPrint conversion cost is only paid for invoices someone
                            # actually wants as a PDF.
                            if st.button("📄 Convert to PDF", key=f"pdf_btn_{key}"):
                                pdf_buf = generate_invoice_pdf(client, invoice_meta, computed)
                                if pdf_buf is None:
                                    st.error("❌ PDF export isn't available - the 'weasyprint' package "
                                              "isn't installed on this server.")
                                else:
                                    st.session_state[f"pdf_bytes_{key}"] = pdf_buf.getvalue()
                            if st.session_state.get(f"pdf_bytes_{key}"):
                                st.download_button(
                                    "⬇️ Download .pdf", data=st.session_state[f"pdf_bytes_{key}"],
                                    file_name=f"{client.get('legal_name','Invoice').replace(' ', '_')}_{invoice_meta['invoice_no'].replace('/', '_')}.pdf",
                                    mime="application/pdf",
                                    key=f"dl_pdf_{key}",
                                )
                        with col_c:
                            send_clicked = st.button("📧 Send by Email", key=f"send_btn_{key}")
                        with col_d:
                            pass

                        if send_clicked:
                            st.session_state[f"show_email_form_{key}"] = True

                        if st.session_state.get(f"show_email_form_{key}"):
                            with st.form(key=f"email_form_{key}"):
                                st.markdown("#### ✉️ Email Settings")
                                smtp_host = st.text_input("SMTP Host", value=st.session_state.baas_smtp_cfg["host"], key=f"smtp_host_{key}")
                                smtp_port = st.number_input("SMTP Port", value=st.session_state.baas_smtp_cfg["port"], key=f"smtp_port_{key}")
                                smtp_user = st.text_input("SMTP Username (sender email)", value=st.session_state.baas_smtp_cfg["username"], key=f"smtp_user_{key}")
                                smtp_pass = st.text_input("SMTP Password / App Password", type="password", key=f"smtp_pass_{key}")
                                to_email = st.text_input("To", value=client.get("contact_email", ""), key=f"to_{key}")
                                cc_email = st.text_input("Cc (comma separated, optional)", key=f"cc_{key}")
                                subject = st.text_input("Subject", value=f"Choice Bank BaaS Invoice - {invoice_meta['invoice_no']}", key=f"subject_{key}")
                                body = st.text_area(
                                    "Body", key=f"body_{key}",
                                    value=(f"Dear {client.get('legal_name','')},\n\n"
                                           f"Please find attached your BaaS invoice {invoice_meta['invoice_no']} "
                                           f"for the period {invoice_meta['period_label']}, "
                                           f"totalling {format_money(computed['net_total'], computed['invoice_currency'])}, "
                                           f"due {invoice_meta['due_date'].strftime('%d/%m/%Y')}.\n\n"
                                           f"Kind regards,\nChoice Bank BaaS Team"),
                                )
                                submitted = st.form_submit_button("Send Now")
                                if submitted:
                                    try:
                                        smtp_cfg = {"host": smtp_host, "port": smtp_port, "username": smtp_user,
                                                    "password": smtp_pass, "use_tls": True}
                                        fname = f"{invoice_meta['invoice_no'].replace('/', '_')}.docx"
                                        send_invoice_email(smtp_cfg, to_email, cc_email, subject, body,
                                                            doc_info["buf"].getvalue(), fname)
                                        baas_db.log_invoice({
                                            "invoice_no": invoice_meta["invoice_no"],
                                            "client_legal_name": client.get("legal_name", ""),
                                            "period": invoice_meta["period_label"],
                                            "currency": computed["invoice_currency"],
                                            "subtotal_gross": computed["subtotal_gross"],
                                            "withholding_tax": computed["total_wht"],
                                            "net_total": computed["net_total"],
                                            "status": "Sent",
                                            "sent_to": to_email,
                                            "sent_at": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                                            "raw_source_text": doc_info["raw_text"],
                                            "generated_by": st.session_state.get("user", {}).get("username", "unknown_user"),
                                        })
                                        st.success(f"✅ Invoice emailed to {to_email}")
                                        st.session_state[f"show_email_form_{key}"] = False
                                    except Exception as e:
                                        st.error(f"❌ Failed to send email: {e}")

    # ------------------------------------------------------------------- Client Directory -----
    with tab_directory:
        st.markdown("### 🗂️ Client / Invoice Directory")
        st.caption(
            "One row per client. 'Match Keywords' (comma separated) is what the parser uses to "
            "recognise a client from the free-text BAAS INCOME description. Banking details are "
            "seeded from the approved templates for Ahadi, Savatech and HongKong FortuneTech - "
            "add the rest here once and every future month reuses them."
        )
        clients_df = baas_db.load_clients()

        with st.expander("➕ Add New Client (seed a new client without editing code)", expanded=False):
            st.caption(
                "Onboard a client here instead of hand-editing a blank row below - this fills in "
                "the same sensible defaults (Choice Microfinance Bank RTGS details, I&M Bank SWIFT "
                "alternative, invoice prefix, match keywords) that a code-seeded client gets. Every "
                "field can still be fine-tuned in the table afterwards, and a client's currency or "
                "banking details can always change again next month by editing that row - this form "
                "is only for creating the client the first time."
            )
            with st.form(key="baas_add_client_form", clear_on_submit=True):
                col1, col2 = st.columns(2)
                with col1:
                    new_legal_name = st.text_input("Legal Name *", key="baas_new_legal_name")
                    new_currency = st.selectbox("Currency *", options=["KES", "USD", "EUR", "GBP"], key="baas_new_currency")
                    new_invoice_group = st.selectbox("Invoice Day *", options=["5th", "10th"], key="baas_new_invoice_group")
                    new_contact_email = st.text_input("Contact Email", key="baas_new_contact_email")
                with col2:
                    new_match_keywords = st.text_input(
                        "Match Keywords (comma separated)", key="baas_new_match_keywords",
                        help="Leave blank to auto-derive from the legal name (same logic the code seed uses)."
                    )
                    new_invoice_prefix = st.text_input(
                        "Invoice Prefix", key="baas_new_invoice_prefix",
                        help="Leave blank to auto-derive from the first word of the legal name."
                    )
                    new_rtgs_account_no = st.text_input("RTGS Account No.", key="baas_new_rtgs_account_no")
                    new_paybill = st.text_input("Paybill (optional)", key="baas_new_paybill")

                st.markdown("**Alternative / SWIFT method**")
                col3, col4 = st.columns(2)
                with col3:
                    new_alt_method_label = st.text_input(
                        "Alt Method Label", value="For SWIFT Transfer;", key="baas_new_alt_method_label"
                    )
                with col4:
                    new_alt_account_no = st.text_input("Alt Account No.", key="baas_new_alt_account_no")

                add_submitted = st.form_submit_button("➕ Add Client", type="primary")
                if add_submitted:
                    legal_name_clean = new_legal_name.strip()
                    if not legal_name_clean:
                        st.error("❌ Legal Name is required.")
                    elif clients_df["legal_name"].str.strip().str.upper().eq(legal_name_clean.upper()).any():
                        st.error(f"❌ A client named '{legal_name_clean}' already exists - edit it in the table below instead.")
                    else:
                        keywords = new_match_keywords.strip() or _default_keywords(legal_name_clean)
                        new_row = _blank_client(legal_name_clean, new_currency, new_invoice_group, keywords)
                        if new_invoice_prefix.strip():
                            new_row["invoice_prefix"] = new_invoice_prefix.strip()
                        new_row["contact_email"] = new_contact_email.strip()
                        new_row["rtgs_account_no"] = new_rtgs_account_no.strip()
                        new_row["paybill"] = new_paybill.strip()
                        new_row["alt_method_label"] = new_alt_method_label.strip() or "For SWIFT Transfer;"
                        new_row["alt_account_no"] = new_alt_account_no.strip()

                        updated_df = pd.concat([clients_df, pd.DataFrame([new_row])], ignore_index=True)
                        baas_db.save_clients(updated_df)
                        st.success(f"✅ Added '{legal_name_clean}'. Fine-tune banking details below if needed.")
                        st.rerun()

        edited_clients = st.data_editor(
            clients_df, use_container_width=True, num_rows="dynamic",
            height=500, key="baas_client_editor",
            column_config={"currency": st.column_config.SelectboxColumn("Currency", options=["KES", "USD", "EUR", "GBP"]),
                            "invoice_group": st.column_config.SelectboxColumn("Invoice Day", options=["5th", "10th"])},
        )
        if st.button("💾 Save Client Directory", type="primary", key="baas_save_clients_btn"):
            baas_db.save_clients(edited_clients)
            st.success("✅ Client directory saved.")
            st.rerun()

    # -------------------------------------------------------------------------- Sent Log ------
    with tab_log:
        st.markdown("### 📨 Sent / Generated Invoices Log")
        log_df = baas_db.load_log()
        if log_df.empty:
            st.info("No invoices generated yet.")
        else:
            st.dataframe(log_df, use_container_width=True)
            csv = log_df.to_csv(index=False).encode('utf-8')
            st.download_button("📥 Download Log (CSV)", data=csv,
                                file_name=f"baas_invoice_log_{datetime.now().strftime('%Y%m%d')}.csv",
                                mime="text/csv")

    # -------------------------------------------------------------------------- Statistics ----
    with tab_stats:
        st.markdown("### 📊 Invoicing Statistics")
        log_df = baas_db.load_log()
        if log_df.empty:
            st.info("No data yet - generate some invoices first.")
        else:
            col1, col2, col3, col4 = st.columns(4)
            col1.metric("🧾 Total Invoices", len(log_df))
            col2.metric("📧 Sent", int((log_df["status"] == "Sent").sum()))
            kes_total = log_df.loc[log_df["currency"] == "KES", "net_total"].astype(float).sum()
            usd_total = log_df.loc[log_df["currency"] == "USD", "net_total"].astype(float).sum()
            col3.metric("💰 KES Net Total", f"KES {kes_total:,.2f}")
            col4.metric("💵 USD Net Total", f"USD {usd_total:,.2f}")

            wht_total = log_df["withholding_tax"].astype(float).sum()
            st.metric("🏛️ Total Withholding Tax Collected (KES)", f"KES {wht_total:,.2f}")

            by_client = log_df.groupby("client_legal_name")["net_total"].apply(lambda s: s.astype(float).sum()).reset_index()
            if not by_client.empty:
                fig = px.bar(by_client, x="client_legal_name", y="net_total", title="Net Invoiced Amount by Client")
                fig.update_layout(xaxis_title="Client", yaxis_title="Net Total")
                st.plotly_chart(fig, use_container_width=True)